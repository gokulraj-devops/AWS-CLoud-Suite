#!/usr/bin/env python3
"""
AWS Cloud Management Suite — Unified Server
Modules: Inventory | Cost Optimization | Security Audit | Monthly Report
Run: python aws_inventory_server.py
Opens: http://localhost:8080
"""
import json
import os
import sys
import subprocess
import shutil
import tempfile
import threading
from datetime import datetime, timedelta
from pathlib import Path
from flask import Flask, request, jsonify, Response, send_file
from flask_cors import CORS
import boto3
from botocore.exceptions import ClientError

app = Flask(__name__)
CORS(app, origins="*")

HTML_PAGE = open(os.path.join(os.path.dirname(os.path.abspath(__file__)), "frontend.html")).read()


@app.route("/", methods=["GET"])
def index():
    return Response(HTML_PAGE, mimetype="text/html")

@app.route("/logo", methods=["GET"])
def serve_logo():
    base = os.path.dirname(os.path.abspath(__file__))
    candidates = [
        ("Shellkode-logo.png", "image/png"), ("shellkode-logo.png", "image/png"),
        ("Shellkode-logo.jpg", "image/jpeg"), ("shellkode-logo.jpg", "image/jpeg"),
        ("logo.png", "image/png"), ("logo.jpg", "image/jpeg"), ("logo.svg", "image/svg+xml"),
        ("shellkode.png", "image/png"),
    ]
    for name, mime in candidates:
        path = os.path.join(base, name)
        if os.path.exists(path):
            resp = send_file(path, mimetype=mime)
            resp.headers["Access-Control-Allow-Origin"] = "*"
            return resp
    return "", 404

# ─── AUTH ─────────────────────────────────────────────────────
@app.route("/auth", methods=["POST"])
def auth():
    b = request.json or {}
    ak, sk = b.get("accessKey","").strip(), b.get("secretKey","").strip()
    if not ak or not sk:
        return jsonify({"success": False, "error": "Credentials required"}), 400
    try:
        sts = boto3.client("sts", aws_access_key_id=ak, aws_secret_access_key=sk, region_name="us-east-1")
        identity = sts.get_caller_identity()
        account_id = identity["Account"]
        arn = identity["Arn"]
        account_name = account_id
        try:
            iam = boto3.client("iam", aws_access_key_id=ak, aws_secret_access_key=sk, region_name="us-east-1")
            aliases = iam.list_account_aliases().get("AccountAliases", [])
            if aliases: account_name = aliases[0]
        except: pass
        return jsonify({"success": True, "accountId": account_id, "accountName": account_name,
                        "userId": identity["UserId"], "arn": arn, "username": arn.split("/")[-1]})
    except ClientError as e:
        code = e.response["Error"]["Code"]
        if code in ("InvalidClientTokenId","SignatureDoesNotMatch","AuthFailure"):
            return jsonify({"success": False, "error": "Invalid credentials."}), 401
        return jsonify({"success": False, "error": str(e)}), 401
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

# ─── REGIONS ──────────────────────────────────────────────────
@app.route("/regions", methods=["POST"])
def list_regions():
    b = request.json or {}
    try:
        ec2 = boto3.client("ec2", aws_access_key_id=b["accessKey"], aws_secret_access_key=b["secretKey"], region_name="us-east-1")
        regions = sorted([r["RegionName"] for r in ec2.describe_regions(AllRegions=False)["Regions"]])
        return jsonify({"success": True, "regions": regions})
    except:
        return jsonify({"success": True, "regions": ["us-east-1","us-east-2","us-west-1","us-west-2",
            "ap-south-1","ap-southeast-1","ap-southeast-2","ap-northeast-1","ap-northeast-2",
            "eu-west-1","eu-west-2","eu-central-1","ca-central-1","sa-east-1"]})

# ─── INVENTORY ────────────────────────────────────────────────
@app.route("/inventory", methods=["POST"])
def inventory():
    b = request.json or {}
    ak, sk, region = b.get("accessKey",""), b.get("secretKey",""), b.get("region","us-east-1")
    services = b.get("services", [])
    s = sess(ak, sk, region)
    result = {}
    for svc in services:
        try:
            fn = COLLECTORS.get(svc)
            result[svc] = fn(s, region) if fn else []
        except Exception as e:
            result[svc] = [{"Error": str(e)}]
    return jsonify({"success": True, "data": result})

# ─── COST OPTIMIZATION ────────────────────────────────────────
@app.route("/cost/report", methods=["POST"])
def cost_report():
    b = request.json or {}
    ak, sk = b.get("accessKey",""), b.get("secretKey","")
    regions_list = b.get("regions", ["us-east-1"])
    try:
        ce = boto3.client("ce", aws_access_key_id=ak, aws_secret_access_key=sk, region_name="us-east-1")
        end = datetime.utcnow().replace(day=1).strftime("%Y-%m-%d")
        start_dt = datetime.utcnow().replace(day=1) - timedelta(days=92)
        start = start_dt.replace(day=1).strftime("%Y-%m-%d")
        # Monthly cost by service
        resp = ce.get_cost_and_usage(
            TimePeriod={"Start": start, "End": end},
            Granularity="MONTHLY",
            Metrics=["UnblendedCost"],
            GroupBy=[{"Type": "DIMENSION", "Key": "SERVICE"}]
        )
        months = []
        svc_monthly = {}
        for r in resp["ResultsByTime"]:
            month = r["TimePeriod"]["Start"][:7]
            if month not in months: months.append(month)
            for grp in r["Groups"]:
                svc = grp["Keys"][0]
                cost = float(grp["Metrics"]["UnblendedCost"]["Amount"])
                if svc not in svc_monthly: svc_monthly[svc] = {}
                svc_monthly[svc][month] = round(cost, 4)
        # Filter out tax and billing-only line items
        EXCLUDE_PATTERNS = ("tax", "support", "premium support", "refund", "credit",
                            "aws marketplace", "marketplace", "savings plan", "discount")
        def _is_billing_noise(svc_name):
            sl = svc_name.lower()
            return any(p in sl for p in EXCLUDE_PATTERNS)

        # Sort by last month cost descending, exclude tax/support
        last_month = months[-1] if months else ""
        sorted_svcs = sorted(svc_monthly.items(), key=lambda x: x[1].get(last_month, 0), reverse=True)
        top_services = [
            {"service": s, "monthly": {m: v.get(m,0) for m in months}}
            for s,v in sorted_svcs[:20]
            if v.get(last_month,0) > 0.01 and not _is_billing_noise(s)
        ]

        # EC2 stopped instances per region
        stopped_ec2 = []
        for region in regions_list:
            try:
                ec2 = boto3.client("ec2", aws_access_key_id=ak, aws_secret_access_key=sk, region_name=region)
                paginator = ec2.get_paginator("describe_instances")
                for page in paginator.paginate(Filters=[{"Name":"instance-state-name","Values":["stopped"]}]):
                    for res in page["Reservations"]:
                        for inst in res["Instances"]:
                            name = next((t["Value"] for t in inst.get("Tags",[]) if t["Key"]=="Name"), "N/A")
                            itype = inst.get("InstanceType","N/A")
                            launch = str(inst.get("LaunchTime",""))[:10]
                            # estimate EBS cost
                            vol_ids = [m["Ebs"]["VolumeId"] for m in inst.get("BlockDeviceMappings",[]) if "Ebs" in m]
                            ebs_gb = 0
                            for vid in vol_ids:
                                try:
                                    vd = ec2.describe_volumes(VolumeIds=[vid])["Volumes"]
                                    ebs_gb += vd[0].get("Size",0) if vd else 0
                                except: pass
                            ebs_cost = round(ebs_gb * 0.10, 2)
                            stopped_ec2.append({"Region":region,"Name":name,"Instance ID":inst["InstanceId"],
                                "Type":itype,"Launch Date":launch,"EBS GB":ebs_gb,"Est. Monthly Savings ($)":ebs_cost})
            except: pass

        # Unattached EBS volumes
        unattached_vols = []
        for region in regions_list:
            try:
                ec2 = boto3.client("ec2", aws_access_key_id=ak, aws_secret_access_key=sk, region_name=region)
                paginator = ec2.get_paginator("describe_volumes")
                for page in paginator.paginate(Filters=[{"Name":"status","Values":["available"]}]):
                    for v in page["Volumes"]:
                        size = v.get("Size",0)
                        cost = round(size * 0.10, 2)
                        name = next((t["Value"] for t in v.get("Tags",[]) if t["Key"]=="Name"), "N/A")
                        unattached_vols.append({"Region":region,"Volume ID":v["VolumeId"],"Name":name,
                            "Size (GB)":size,"Type":v.get("VolumeType","N/A"),
                            "Created":str(v.get("CreateTime",""))[:10],"Est. Monthly Savings ($)":cost})
            except: pass

        # Unattached EIPs
        unattached_eips = []
        for region in regions_list:
            try:
                ec2 = boto3.client("ec2", aws_access_key_id=ak, aws_secret_access_key=sk, region_name=region)
                resp2 = ec2.describe_addresses()
                for a in resp2["Addresses"]:
                    if "AssociationId" not in a:
                        unattached_eips.append({"Region":region,"Allocation ID":a.get("AllocationId","N/A"),
                            "Public IP":a.get("PublicIp","N/A"),"Est. Monthly Savings ($)":3.65})
            except: pass

        # Old snapshots (>90 days)
        old_snapshots = []
        for region in regions_list:
            try:
                ec2 = boto3.client("ec2", aws_access_key_id=ak, aws_secret_access_key=sk, region_name=region)
                sts2 = boto3.client("sts", aws_access_key_id=ak, aws_secret_access_key=sk)
                owner_id = sts2.get_caller_identity()["Account"]
                paginator = ec2.get_paginator("describe_snapshots")
                cutoff = datetime.utcnow() - timedelta(days=90)
                for page in paginator.paginate(OwnerIds=[owner_id]):
                    for snap in page["Snapshots"]:
                        start_time = snap.get("StartTime")
                        if start_time:
                            if hasattr(start_time, "replace"):
                                start_naive = start_time.replace(tzinfo=None)
                            else:
                                start_naive = datetime.strptime(str(start_time)[:19], "%Y-%m-%dT%H:%M:%S")
                            if start_naive < cutoff:
                                size = snap.get("VolumeSize",0)
                                cost = round(size * 0.05, 2)
                                old_snapshots.append({"Region":region,"Snapshot ID":snap["SnapshotId"],
                                    "Volume ID":snap.get("VolumeId","N/A"),"Size (GB)":size,
                                    "Created":str(snap.get("StartTime",""))[:10],
                                    "Description":snap.get("Description","N/A")[:50],
                                    "Est. Monthly Savings ($)":cost})
            except: pass

        # Old / unused AMIs (>90 days, owned by this account)
        old_amis = []
        _ami_owner_id = ""
        try:
            _ami_owner_id = boto3.client("sts", aws_access_key_id=ak,
                aws_secret_access_key=sk).get_caller_identity()["Account"]
        except: _ami_owner_id = "self"
        for region in regions_list:
            try:
                ec2 = boto3.client("ec2", aws_access_key_id=ak, aws_secret_access_key=sk, region_name=region)
                images = ec2.describe_images(Owners=[_ami_owner_id or "self"])["Images"]
                ami_cutoff = datetime.utcnow() - timedelta(days=90)
                for img in images:
                    cdate_str = img.get("CreationDate", "")
                    if not cdate_str: continue
                    try:
                        cdt = datetime.strptime(cdate_str[:19], "%Y-%m-%dT%H:%M:%S")
                    except: continue
                    if cdt >= ami_cutoff: continue
                    snap_ids = [bdm["Ebs"]["SnapshotId"] for bdm in img.get("BlockDeviceMappings",[]) if "Ebs" in bdm and "SnapshotId" in bdm.get("Ebs",{})]
                    snap_size_gb = 0
                    for sid in snap_ids:
                        try:
                            snaps = ec2.describe_snapshots(SnapshotIds=[sid])["Snapshots"]
                            snap_size_gb += snaps[0].get("VolumeSize",0) if snaps else 0
                        except: pass
                    ami_cost = round(snap_size_gb * 0.05, 2)
                    old_amis.append({
                        "Region": region, "AMI ID": img["ImageId"],
                        "Name": img.get("Name","N/A")[:60],
                        "State": img.get("State","N/A"),
                        "Architecture": img.get("Architecture","N/A"),
                        "Creation Date": cdate_str[:10],
                        "Snapshot Count": len(snap_ids),
                        "Snapshot Size (GB)": snap_size_gb,
                        "Est. Monthly Savings ($)": ami_cost,
                    })
            except: pass

        # Savings summary table
        def _total_savings(rows):
            return round(sum(float(r.get("Est. Monthly Savings ($)",0) or 0) for r in rows), 2)
        savings_summary = [
            {"Resource": "Stopped EC2 Instances",    "Count": len(stopped_ec2),
             "Monthly Savings ($)": _total_savings(stopped_ec2),
             "Yearly Savings ($)":  round(_total_savings(stopped_ec2)*12, 2),
             "Note": "EBS storage cost while instance is stopped"},
            {"Resource": "Unattached EBS Volumes",   "Count": len(unattached_vols),
             "Monthly Savings ($)": _total_savings(unattached_vols),
             "Yearly Savings ($)":  round(_total_savings(unattached_vols)*12, 2),
             "Note": "Volumes not attached to any instance"},
            {"Resource": "Unattached Elastic IPs",   "Count": len(unattached_eips),
             "Monthly Savings ($)": _total_savings(unattached_eips),
             "Yearly Savings ($)":  round(_total_savings(unattached_eips)*12, 2),
             "Note": "$3.65/IP/month when not associated"},
            {"Resource": "Old EBS Snapshots (>90d)", "Count": len(old_snapshots),
             "Monthly Savings ($)": _total_savings(old_snapshots),
             "Yearly Savings ($)":  round(_total_savings(old_snapshots)*12, 2),
             "Note": "Snapshots older than 90 days"},
            {"Resource": "Old AMIs (>90d)",          "Count": len(old_amis),
             "Monthly Savings ($)": _total_savings(old_amis),
             "Yearly Savings ($)":  round(_total_savings(old_amis)*12, 2),
             "Note": "Old AMIs with associated EBS snapshots"},
        ]
        total_m = round(sum(r["Monthly Savings ($)"] for r in savings_summary), 2)
        savings_summary.append({"Resource": "TOTAL POTENTIAL SAVINGS", "Count": "",
            "Monthly Savings ($)": total_m, "Yearly Savings ($)": round(total_m*12, 2), "Note": ""})

        return jsonify({"success": True, "months": months, "topServices": top_services,
            "stoppedEC2": stopped_ec2, "unattachedVolumes": unattached_vols,
            "unattachedEIPs": unattached_eips, "oldSnapshots": old_snapshots,
            "oldAMIs": old_amis, "savingsSummary": savings_summary})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

# ─── SECURITY AUDIT ───────────────────────────────────────────
# ── Prowler helpers ──────────────────────────────────────────
def _find_prowler():
    """Return path to prowler executable: checks venv first, then PATH."""
    script_dir = Path(os.path.dirname(os.path.abspath(__file__)))
    candidates = [
        script_dir / "venv" / "bin" / "prowler",
        script_dir / ".venv" / "bin" / "prowler",
        Path(sys.executable).parent / "prowler",
    ]
    for c in candidates:
        if c.exists():
            return str(c)
    found = shutil.which("prowler")
    return found  # None if not found

def _install_prowler():
    """Install prowler into the same Python environment as this server."""
    pip = str(Path(sys.executable).parent / "pip")
    if not os.path.exists(pip):
        pip = shutil.which("pip") or shutil.which("pip3") or f"{sys.executable} -m pip"
    result = subprocess.run(
        [sys.executable, "-m", "pip", "install", "prowler", "--quiet"],
        capture_output=True, text=True, timeout=300
    )
    return result.returncode == 0, result.stdout + result.stderr

# Global scan state for SSE streaming
_scan_state = {"status": "idle", "message": "", "csvPath": "", "error": ""}

@app.route("/security/check", methods=["POST"])
def security_check():
    """Check if prowler is installed and optionally install it."""
    b = request.json or {}
    do_install = b.get("install", False)
    prowler_path = _find_prowler()
    if prowler_path:
        return jsonify({"installed": True, "path": prowler_path})
    if do_install:
        ok, log = _install_prowler()
        if ok:
            prowler_path = _find_prowler()
            if prowler_path:
                return jsonify({"installed": True, "path": prowler_path, "justInstalled": True})
        return jsonify({"installed": False, "error": f"Auto-install failed: {log[-300:]}"})
    return jsonify({"installed": False, "error": "Prowler not found. Click 'Auto-Install & Scan'."})

@app.route("/security/run", methods=["POST"])
def security_run():
    b = request.json or {}
    ak, sk = b.get("accessKey",""), b.get("secretKey","")
    session_token = b.get("sessionToken", "")

    # Resolve prowler path
    prowler_path = _find_prowler()
    if not prowler_path:
        ok, log = _install_prowler()
        prowler_path = _find_prowler() if ok else None
        if not prowler_path:
            return jsonify({"success": False,
                "error": "Prowler is not installed. Run: pip install prowler (inside venv)"})

    # Detect prowler version — v5 outputs to stderr, so check both
    try:
        ver_proc = subprocess.run([prowler_path, "--version"],
            capture_output=True, text=True, timeout=15)
        ver_out = (ver_proc.stdout + ver_proc.stderr).strip()
        # e.g. "prowler 5.22.0", "prowler 4.3.2", or bare "3.15.0"
        ver_str = ver_out.split()[-1] if ver_out else "5"
        prowler_major = int(ver_str.split(".")[0])
        if prowler_major not in (3, 4, 5):
            prowler_major = 5
    except Exception:
        prowler_major = 5  # assume latest

    script_dir = Path(os.path.dirname(os.path.abspath(__file__)))
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = script_dir / f"prowler_output_{timestamp}"
    output_dir.mkdir(parents=True, exist_ok=True)

    env = os.environ.copy()
    env.update({
        "AWS_ACCESS_KEY_ID": ak,
        "AWS_SECRET_ACCESS_KEY": sk,
        "AWS_DEFAULT_REGION": "us-east-1",
    })
    if session_token:
        env["AWS_SESSION_TOKEN"] = session_token

    # ── Detect enabled AWS regions for this account ──
    # Passing only enabled regions to Prowler cuts scan time dramatically.
    # Without this, Prowler iterates all ~30 AWS regions even if only 2 are used.
    enabled_regions = []
    try:
        import boto3
        ec2_client = boto3.client(
            "ec2",
            aws_access_key_id=ak,
            aws_secret_access_key=sk,
            aws_session_token=session_token or None,
            region_name="us-east-1"
        )
        resp = ec2_client.describe_regions(AllRegions=False)
        enabled_regions = sorted([r["RegionName"] for r in resp["Regions"]])
    except Exception:
        enabled_regions = []  # let Prowler scan all if we can't detect

    # ── Build command ──
    # Service names differ between Prowler v4 and v5 (e.g. lambda -> awslambda).
    # We auto-discover valid names by running --list-services, then intersect
    # with our desired set so we always pass only valid names.

    # Desired services — covers the most important security checks
    DESIRED_SERVICES = {
        # v4 name      v5 name
        "iam",         # iam         (unchanged)
        "s3",          # s3          (unchanged)
        "ec2",         # ec2         (unchanged)
        "rds",         # rds         (unchanged)
        "cloudtrail",  # cloudtrail  (unchanged)
        "cloudwatch",  # cloudwatch  (unchanged)
        "config",      # config      (unchanged)
        "guardduty",   # guardduty   (unchanged)
        "kms",         # kms         (unchanged)
        "lambda",      # v4 name
        "awslambda",   # v5 name
        "eks",         # eks         (unchanged)
        "secretsmanager",
        "vpc",
        "elbv2",
        "sqs",
        "acm",
        "ecr",
    }

    # Auto-discover valid service names for this exact Prowler version
    valid_services = set()
    try:
        ls_proc = subprocess.run(
            [prowler_path, "aws", "--list-services"],
            capture_output=True, text=True, timeout=30, env=env
        )
        ls_out = ls_proc.stdout + ls_proc.stderr
        # Output lines look like:  "  - awslambda" or "awslambda"
        import re
        for line in ls_out.splitlines():
            m = re.search(r'([a-z][a-z0-9_]+)', line.strip())
            if m:
                valid_services.add(m.group(1))
    except Exception:
        pass  # fall through to use desired set as-is

    if valid_services:
        # Use only services that both exist in Prowler and we want
        services_to_run = sorted(DESIRED_SERVICES & valid_services)
        if not services_to_run:
            # Fallback: run without service filter (full scan)
            services_to_run = None
    else:
        # Could not list services — build best-guess list based on version
        if prowler_major >= 5:
            services_to_run = [
                "iam", "s3", "ec2", "rds", "cloudtrail", "cloudwatch",
                "config", "guardduty", "kms", "awslambda", "eks",
                "secretsmanager", "vpc", "elbv2", "sqs", "acm", "ecr",
            ]
        else:
            services_to_run = [
                "iam", "s3", "ec2", "rds", "cloudtrail", "cloudwatch",
                "config", "guardduty", "kms", "lambda", "eks",
                "secretsmanager", "vpc", "elbv2", "sqs", "acm", "ecr",
            ]

    # Build --region flags: Prowler v4/v5 uses "-f region1 -f region2" (repeatable)
    # This limits scanning to only your enabled regions — massive speed improvement
    region_flags = []
    for r in enabled_regions:
        region_flags += ["-f", r]

    if prowler_major >= 4:
        cmd = [prowler_path, "aws",
               "--output-formats", "csv", "html",
               "--output-directory", str(output_dir),
               "--ignore-exit-code-3"]
        if services_to_run:
            cmd += ["--services"] + services_to_run
        if region_flags:
            cmd += region_flags
    else:
        # v3 — no --services flag, uses same -f flag for regions
        cmd = [prowler_path,
               "-M", "csv,html",
               "-o", str(output_dir),
               "--ignore-exit-code-3"]
        if region_flags:
            cmd += region_flags

    try:
        result = subprocess.run(
            cmd, env=env, capture_output=True, text=True, timeout=3600
        )

        # Find the main CSV findings file (largest non-compliance CSV)
        csv_files = list(output_dir.rglob("*.csv"))
        main_csv = None
        for f in sorted(csv_files, key=lambda x: x.stat().st_size, reverse=True):
            nl = f.name.lower()
            if "compliance" not in nl and "overview" not in nl:
                main_csv = f; break
        if not main_csv and csv_files:
            main_csv = max(csv_files, key=lambda x: x.stat().st_size)

        html_files = list(output_dir.rglob("*.html"))
        main_html = html_files[0] if html_files else None

        if main_csv:
            return jsonify({
                "success": True,
                "csvPath": str(main_csv),
                "htmlPath": str(main_html) if main_html else "",
                "outputDir": str(output_dir),
                "prowlerVersion": prowler_major,
                "regionsScanned": enabled_regions,
                "servicesScanned": services_to_run or [],
            })

        stderr_tail = (result.stdout + result.stderr)[-1200:]
        return jsonify({"success": False,
            "error": f"Prowler ran (exit {result.returncode}) but no CSV was generated.",
            "detail": stderr_tail})
    except FileNotFoundError:
        return jsonify({"success": False, "error": f"Prowler not found at: {prowler_path}"})
    except subprocess.TimeoutExpired:
        return jsonify({"success": False, "error": "Prowler scan timed out (60 min limit)."})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


# ── Column name aliases: Prowler v3 / v4 / v5 → normalised ──────────────────
# Prowler v5 changed many column names. All known variants are mapped here.
_COL_ALIASES = {
    # ── v3 (semicolon CSV) ───────────────────────────────────────────────────
    "SERVICE_NAME":             "SERVICE_NAME",
    "CHECK_TITLE":              "CHECK_TITLE",
    "STATUS":                   "STATUS",
    "REGION":                   "REGION",
    "RESOURCE_ARN":             "RESOURCE_ARN",
    "RESOURCE_ID":              "RESOURCE_ID",
    "STATUS_EXTENDED":          "STATUS_EXTENDED",
    "SEVERITY":                 "SEVERITY",
    # ── v4 ───────────────────────────────────────────────────────────────────
    "SERVICENAME":              "SERVICE_NAME",
    "CHECKID":                  "CHECK_TITLE",
    "CHECKTITLE":               "CHECK_TITLE",
    "CHECK_ID":                 "CHECK_TITLE",
    "CHECK_TYPE":               "CHECK_TITLE",
    "FINDING_STATUS":           "STATUS",
    "STATUS_CODE":              "STATUS",
    "RESOURCE_UID":             "RESOURCE_ID",
    "RESOURCEID":               "RESOURCE_ID",
    "RESOURCEARN":              "RESOURCE_ARN",
    "RESOURCE_TYPE":            "RESOURCE_ARN",
    "EXTENDED_STATUS":          "STATUS_EXTENDED",
    "FINDING_DESCRIPTION":      "STATUS_EXTENDED",
    "RISK":                     "SEVERITY",
    "SEVERITY_LEVEL":           "SEVERITY",
    # ── v5 new column names ──────────────────────────────────────────────────
    "CHECKMETADATA_SERVICE":    "SERVICE_NAME",
    "CHECKMETADATA_CHECKTITLE": "CHECK_TITLE",
    "CHECKMETADATA_SEVERITY":   "SEVERITY",
    "FINDING_UID":              "RESOURCE_ID",
    "RESOURCE_NAME":            "RESOURCE_ID",
    "DESCRIPTION":              "STATUS_EXTENDED",
    "RISK_DETAILS":             "STATUS_EXTENDED",
    "STATUS_DETAIL":            "STATUS_EXTENDED",
    "RESOURCE_DETAILS":         "STATUS_EXTENDED",
    # ── lowercase variants (v5 sometimes writes lowercase headers) ────────────
    "service_name":             "SERVICE_NAME",
    "check_title":              "CHECK_TITLE",
    "status":                   "STATUS",
    "region":                   "REGION",
    "resource_arn":             "RESOURCE_ARN",
    "resource_id":              "RESOURCE_ID",
    "status_extended":          "STATUS_EXTENDED",
    "severity":                 "SEVERITY",
    "finding_status":           "STATUS",
    "resource_uid":             "RESOURCE_ID",
    "check_id":                 "CHECK_TITLE",
    "finding_description":      "STATUS_EXTENDED",
    "risk":                     "SEVERITY",
    "resource_name":            "RESOURCE_ID",
    "finding_uid":              "RESOURCE_ID",
    "description":              "STATUS_EXTENDED",
    "status_detail":            "STATUS_EXTENDED",
    "checkmetadata_service":    "SERVICE_NAME",
    "checkmetadata_checktitle": "CHECK_TITLE",
    "checkmetadata_severity":   "SEVERITY",
}


def _to_series(obj, name="col"):
    """
    Guarantee we always get a 1-D pandas Series, never a DataFrame.
    When a CSV has duplicate column names, df['col'] returns a DataFrame —
    this collapses it to the first column and converts to string.
    """
    import pandas as pd
    if isinstance(obj, pd.DataFrame):
        obj = obj.iloc[:, 0]
    return obj.fillna("").astype(str)


def _safe_series(df, col):
    """
    Safely extract column `col` from df as a flat 1-D string Series.
    Guards against: duplicate column names (returns DataFrame), MultiIndex,
    object arrays containing sub-DataFrames. Always returns a Series of strings.
    """
    import pandas as pd
    val = df[col]
    # If indexing returned a DataFrame (duplicate column names), take first column
    if isinstance(val, pd.DataFrame):
        val = val.iloc[:, 0]
    # If it's somehow still not 1-D, flatten
    if hasattr(val, "ndim") and val.ndim > 1:
        val = val.iloc[:, 0]
    return val.fillna("").astype(str).replace("nan", "").replace("None", "")


def _normalise_df(df):
    """
    Rename Prowler v3 / v4 / v5 columns → unified schema.
    Fully guards against:
      - Duplicate column names (pandas returns DataFrame on index → .str fails)
      - MultiIndex column headers
      - Multiple source columns renaming to the same target (creates new dupes)
      - Mixed-case / lowercase column headers
      - v5 PASSED/FAILED status values
    """
    import pandas as pd

    # ── Step 1: Flatten MultiIndex columns ──────────────────────────────────
    if isinstance(df.columns, pd.MultiIndex):
        df.columns = ["_".join(str(c) for c in col).strip() for col in df.columns]

    # ── Step 2: Strip whitespace from column names ───────────────────────────
    df.columns = [str(c).strip() for c in df.columns]

    # ── Step 3: Drop duplicate column names (keep first occurrence) ──────────
    # Root cause of "DataFrame object has no attribute 'str'":
    # Prowler v5 CSV sometimes repeats headers; pandas returns a DataFrame
    # when you do df["DUPLICATE_COL"], which has no .str accessor.
    df = df.loc[:, ~df.columns.duplicated(keep="first")]

    # ── Step 4: Force every column to a flat 1-D string Series ──────────────
    for col in list(df.columns):
        df[col] = _to_series(df[col])

    # ── Step 5: Rename columns to normalised names ───────────────────────────
    # Build rename map: first match wins (handles both upper and lower variants)
    rename_map = {}
    for col in list(df.columns):
        # Try exact match, then UPPER match
        norm = _COL_ALIASES.get(col) or _COL_ALIASES.get(col.upper())
        if norm and col != norm and col not in rename_map:
            rename_map[col] = norm

    # Apply rename — this may create NEW duplicate targets
    # e.g. both FINDING_STATUS and STATUS_CODE map to STATUS
    df = df.rename(columns=rename_map)

    # ── Step 6: Drop duplicates AGAIN after rename (keep first) ─────────────
    df = df.loc[:, ~df.columns.duplicated(keep="first")]

    # ── Step 7: Force all columns to flat Series once more (post-rename) ─────
    for col in list(df.columns):
        df[col] = _to_series(df[col])

    # ── Step 8: Ensure all required columns exist ────────────────────────────
    REQUIRED = ["STATUS", "SERVICE_NAME", "CHECK_TITLE", "SEVERITY",
                "REGION", "RESOURCE_ARN", "RESOURCE_ID", "STATUS_EXTENDED"]
    for req in REQUIRED:
        if req not in df.columns:
            df[req] = pd.Series([""] * len(df), dtype=str)
        # Always re-flatten in case rename produced a stale DataFrame reference
        df[req] = _safe_series(df, req)

    # ── Step 9: Normalise STATUS ──────────────────────────────────────────────
    # v3/v4 : PASS / FAIL / WARNING / MUTED / MANUAL
    # v5    : PASSED / FAILED / MUTED / MANUAL
    STATUS_MAP = {
        "PASS":    "PASS",
        "PASSED":  "PASS",
        "FAIL":    "FAIL",
        "FAILED":  "FAIL",
        "WARNING": "FAIL",   # v3 WARNING = finding = treat as FAIL
        "MUTED":   "PASS",   # suppressed finding
        "MANUAL":  "PASS",   # manual check, not automated failure
    }
    status_series = _safe_series(df, "STATUS")
    df["STATUS"] = status_series.str.strip().str.upper().map(
        lambda x: STATUS_MAP.get(x, "FAIL" if x not in ("", "nan") else "PASS")
    )

    # ── Step 10: Normalise SEVERITY to lowercase ──────────────────────────────
    df["SEVERITY"] = _safe_series(df, "SEVERITY").str.strip().str.lower()

    # ── Step 11: Normalise SERVICE_NAME — fill blanks ─────────────────────────
    svc = _safe_series(df, "SERVICE_NAME").str.strip()
    svc[svc.isin(["", "nan", "None", "none"])] = "Other"
    df["SERVICE_NAME"] = svc

    return df


@app.route("/security/format", methods=["POST"])
def security_format():
    """Parse CSV → JSON for UI rendering AND generate styled XLSX server-side."""
    b = request.json or {}
    csv_path = b.get("csvPath", "")
    cust_name = b.get("customerName", "Report")
    account_id = b.get("accountId", "")

    if not csv_path or not os.path.exists(csv_path):
        return jsonify({"success": False, "error": "CSV file not found"})
    try:
        import pandas as pd
        import numpy as np
        import xlsxwriter
        import warnings
        warnings.simplefilter(action='ignore', category=pd.errors.DtypeWarning)

        # Try common separators — Prowler v3 uses ';', v4/v5 use ','
        # Also try tab and pipe in case of version differences
        df = None
        for sep in [",", ";", "	", "|"]:
            try:
                _df = pd.read_csv(csv_path, encoding="utf-8", sep=sep,
                                  on_bad_lines="skip", low_memory=False)
                if len(_df.columns) > 3:
                    df = _df; break
            except Exception:
                try:
                    _df = pd.read_csv(csv_path, encoding="ISO-8859-1", sep=sep,
                                      on_bad_lines="skip", low_memory=False)
                    if len(_df.columns) > 3:
                        df = _df; break
                except Exception:
                    continue
        if df is None:
            return jsonify({"success": False, "error": "Could not parse CSV file"})

        # ── Deduplicate columns FIRST (before any .str ops) ──────────────────
        # Prowler v5 CSV sometimes has repeated column headers; pandas returns
        # a DataFrame (not Series) when you index a duplicate column name,
        # which causes "DataFrame object has no attribute 'str'".
        # NOTE: _normalise_df also does this — doing it here too so the loop
        # below always operates on guaranteed 1-D Series.
        df.columns = [str(c).strip() for c in df.columns]
        df = df.loc[:, ~df.columns.duplicated(keep="first")]

        # Force every column to a flat 1-D string Series
        for col in list(df.columns):
            df[col] = _to_series(df[col])

        # Normalise column names + STATUS values across all Prowler versions
        df = _normalise_df(df)

        if "STATUS" not in df.columns:
            return jsonify({"success": False,
                "error": f"STATUS column not found after normalisation. Raw columns: {list(df.columns)}"})

        df_fail = df[df["STATUS"] == "FAIL"].copy()
        total_checks = len(df)
        total_fails  = len(df_fail)

        # Build summary + details for UI
        summary = []
        details = {}
        sev_order = {"critical": 0, "high": 1, "medium": 2, "low": 3, "informational": 4}

        for svc, grp in df_fail.groupby("SERVICE_NAME"):
            svc = str(svc).strip() or "Other"
            grp = grp.copy()
            # Re-flatten any columns that may have become DataFrames inside group
            for _c in list(grp.columns):
                grp[_c] = _to_series(grp[_c])

            _tmp = grp[["CHECK_TITLE", "SEVERITY"]].drop_duplicates().copy()
            _tmp["_sev_ord"] = _to_series(_tmp["SEVERITY"]).str.lower().map(
                lambda x: sev_order.get(x, 9))
            unique_checks = _tmp.sort_values("_sev_ord").drop(columns=["_sev_ord"])
            checks = unique_checks.to_dict("records")
            summary.append({"service": svc, "checks": checks, "count": len(grp)})
            detail_cols = ["CHECK_TITLE", "STATUS", "SERVICE_NAME", "SEVERITY",
                           "RESOURCE_ARN", "REGION", "STATUS_EXTENDED", "RESOURCE_ID"]
            detail_cols = [c for c in detail_cols if c in grp.columns]
            rows = (grp[detail_cols]
                    .fillna("")
                    .replace([np.inf, -np.inf], "")
                    .to_dict("records"))
            details[svc] = rows

        # ── Generate styled XLSX via prowler_formatter.py ──────────────────
        output_dir_path = Path(csv_path).parent
        safe_cust = "".join(c if c.isalnum() or c in "-_" else "_" for c in cust_name)
        xlsx_filename = (f"{safe_cust}_{account_id}_security_audit.xlsx"
                         if account_id else f"{safe_cust}_security_audit.xlsx")
        xlsx_path = str(output_dir_path / xlsx_filename)

        script_dir_fmt = Path(os.path.dirname(os.path.abspath(__file__)))
        formatter_script = script_dir_fmt / "prowler_formatter.py"
        fmt_ok = False
        if formatter_script.exists():
            try:
                fmt_proc = subprocess.run(
                    [sys.executable, str(formatter_script),
                     "--input",    csv_path,
                     "--output",   xlsx_path,
                     "--customer", cust_name,
                     "--account",  account_id or ""],
                    capture_output=True, text=True, timeout=180
                )
                fmt_ok = fmt_proc.returncode == 0 and os.path.exists(xlsx_path)
            except Exception:
                fmt_ok = False

        # Fallback: inline generation if prowler_formatter.py not found or failed
        if not fmt_ok:
            now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            SEV_COLOURS = {
                "critical":      {"bg": "C62828", "fg": "FFFFFF"},
                "high":          {"bg": "EF6C00", "fg": "FFFFFF"},
                "medium":        {"bg": "F9A825", "fg": "000000"},
                "low":           {"bg": "2E7D32", "fg": "FFFFFF"},
                "informational": {"bg": "1565C0", "fg": "FFFFFF"},
            }
            with xlsxwriter.Workbook(xlsx_path) as wb:
                hdr_fmt = wb.add_format({
                    "bold": True, "font_name": "Arial", "font_size": 11,
                    "bg_color": "#1A3A6B", "font_color": "#FFFFFF",
                    "border": 1, "align": "center", "valign": "vcenter", "text_wrap": True,
                })
                title_fmt = wb.add_format({
                    "bold": True, "font_name": "Arial", "font_size": 14,
                    "bg_color": "#0C1929", "font_color": "#FFFFFF",
                    "border": 1, "align": "left", "valign": "vcenter",
                })
                label_fmt = wb.add_format({
                    "bold": True, "font_name": "Arial", "font_size": 10,
                    "bg_color": "#112238", "font_color": "#AEC9F0",
                    "border": 1, "align": "left", "valign": "vcenter",
                })
                value_fmt = wb.add_format({
                    "font_name": "Arial", "font_size": 10,
                    "bg_color": "#162B45", "font_color": "#E2E8F0",
                    "border": 1, "align": "left", "valign": "vcenter",
                })
                cell_odd  = wb.add_format({
                    "font_name": "Arial", "font_size": 10,
                    "bg_color": "#FFFFFF", "font_color": "#1E293B",
                    "border": 1, "align": "left", "valign": "vcenter",
                })
                cell_even = wb.add_format({
                    "font_name": "Arial", "font_size": 10,
                    "bg_color": "#F0F4F8", "font_color": "#1E293B",
                    "border": 1, "align": "left", "valign": "vcenter",
                })
                remark_fmt = wb.add_format({
                    "font_name": "Arial", "font_size": 10,
                    "bg_color": "#FFFDE7", "font_color": "#33691E",
                    "border": 1, "align": "left", "valign": "vcenter",
                })
                merge_fmt = wb.add_format({
                    "bold": True, "font_name": "Arial", "font_size": 10,
                    "bg_color": "#E8EFF8", "font_color": "#1E293B",
                    "border": 1, "align": "center", "valign": "vcenter", "text_wrap": True,
                })
    
                # Pre-build severity format cache to avoid xlsxwriter limit
                _sev_fmt_cache = {}
                def sev_fmt(sev_str):
                    key = (sev_str or "").lower()
                    if key not in _sev_fmt_cache:
                        c = SEV_COLOURS.get(key, {"bg": "475569", "fg": "FFFFFF"})
                        _sev_fmt_cache[key] = wb.add_format({
                            "bold": True, "font_name": "Arial", "font_size": 10,
                            "bg_color": f"#{c['bg']}", "font_color": f"#{c['fg']}",
                            "border": 1, "align": "center", "valign": "vcenter",
                        })
                    return _sev_fmt_cache[key]
    
                fail_fmt = wb.add_format({
                    "bold": True, "font_name": "Arial", "font_size": 10,
                    "bg_color": "#C62828", "font_color": "#FFFFFF",
                    "border": 1, "align": "center", "valign": "vcenter",
                })
    
                # ══ SUMMARY SHEET ══
                ws = wb.add_worksheet("Summary")
                ws.set_column(0, 0, 28)
                ws.set_column(1, 1, 60)
                ws.set_column(2, 2, 16)
                ws.set_row(0, 32)
                ws.merge_range(0, 0, 0, 2,
                    f"Security Audit Report — {cust_name}  |  Account: {account_id}  |  {now_str}",
                    title_fmt)
    
                meta = [
                    ("Total Checks Run",  str(total_checks)),
                    ("Total Failures",    str(total_fails)),
                    ("Services Affected", str(len(summary))),
                    ("Prowler Output",    csv_path),
                ]
                for i, (k, v) in enumerate(meta):
                    ws.write(i + 1, 0, k, label_fmt)
                    ws.merge_range(i + 1, 1, i + 1, 2, v, value_fmt)
    
                hdr_row = len(meta) + 2
                ws.set_row(hdr_row, 20)
                ws.write(hdr_row, 0, "SERVICE NAME", hdr_fmt)
                ws.write(hdr_row, 1, "CHECK TITLE",  hdr_fmt)
                ws.write(hdr_row, 2, "SEVERITY",     hdr_fmt)
    
                data_row = hdr_row + 1
                for item in summary:
                    svc    = item["service"]
                    checks = item["checks"]
                    n      = len(checks)
                    if n > 1:
                        ws.merge_range(data_row, 0, data_row + n - 1, 0, svc, merge_fmt)
                    else:
                        ws.write(data_row, 0, svc, merge_fmt)
                    for chk in checks:
                        rf = cell_odd if (data_row % 2 == 0) else cell_even
                        ws.write(data_row, 1, chk.get("CHECK_TITLE", ""), rf)
                        ws.write(data_row, 2, chk.get("SEVERITY", ""),
                                 sev_fmt(chk.get("SEVERITY", "")))
                        data_row += 1
    
                # ══ PER-SERVICE SHEETS with Remark column + autofilter ══
                detail_cols_order = ["CHECK_TITLE", "SEVERITY", "STATUS",
                                     "RESOURCE_ID", "RESOURCE_ARN",
                                     "REGION", "SERVICE_NAME", "STATUS_EXTENDED", "Remark"]
                svc_hdr_colours = [
                    "#1A3A6B","#243C72","#1C3560","#2A4880",
                    "#16305A","#0F2850","#2E4A7A","#1F3A5F",
                ]
                col_widths = {
                    "CHECK_TITLE": 55, "SEVERITY": 14, "STATUS": 10,
                    "RESOURCE_ID": 32, "RESOURCE_ARN": 55,
                    "REGION": 18, "SERVICE_NAME": 22, "STATUS_EXTENDED": 60, "Remark": 30,
                }
    
                for sh_idx, item in enumerate(summary):
                    svc       = item["service"]
                    rows_data = details.get(svc, [])
                    if not rows_data:
                        continue
                    sh_colour = svc_hdr_colours[sh_idx % len(svc_hdr_colours)]
                    sws = wb.add_worksheet(svc[:31])
    
                    sh_hdr = wb.add_format({
                        "bold": True, "font_name": "Arial", "font_size": 10,
                        "bg_color": sh_colour, "font_color": "#FFFFFF",
                        "border": 1, "align": "center", "valign": "vcenter", "text_wrap": True,
                    })
                    sh_title = wb.add_format({
                        "bold": True, "font_name": "Arial", "font_size": 12,
                        "bg_color": "#0C1929", "font_color": "#FFFFFF",
                        "border": 1, "align": "left", "valign": "vcenter",
                    })
    
                    # Columns present in data + Remark
                    data_keys = set(rows_data[0].keys()) if rows_data else set()
                    avail = [c for c in detail_cols_order
                             if c in data_keys or c == "Remark"]
                    if not avail:
                        avail = list(data_keys) + ["Remark"]
    
                    n_cols = len(avail)
                    sws.merge_range(0, 0, 0, n_cols - 1,
                        f"{svc}  —  {len(rows_data)} failures  |  Account: {account_id}",
                        sh_title)
                    sws.set_row(0, 24)
                    sws.set_row(1, 20)
                    for ci, col in enumerate(avail):
                        sws.write(1, ci, col.replace("_", " "), sh_hdr)
    
                    # Autofilter on header row
                    sws.autofilter(1, 0, 1 + len(rows_data), n_cols - 1)
    
                    for ri, row in enumerate(rows_data):
                        rf = cell_odd if ri % 2 == 0 else cell_even
                        sws.set_row(ri + 2, 16)
                        for ci, col in enumerate(avail):
                            val = row.get(col, "")
                            try:
                                if pd.isna(val): val = ""
                            except Exception:
                                pass
                            if col == "Remark":
                                sws.write(ri + 2, ci, "", remark_fmt)
                            elif col == "SEVERITY":
                                sws.write(ri + 2, ci, str(val), sev_fmt(str(val)))
                            elif col == "STATUS":
                                sws.write(ri + 2, ci, str(val), fail_fmt)
                            else:
                                sws.write(ri + 2, ci, str(val), rf)
    
                    for ci, col in enumerate(avail):
                        sws.set_column(ci, ci, col_widths.get(col, 22))
    
        return jsonify({
            "success": True,
            "summary": summary,
            "details": details,
            "totalFails": total_fails,
            "totalChecks": total_checks,
            "totalServices": len(summary),
            "xlsxPath": xlsx_path,
            "xlsxFilename": os.path.basename(xlsx_path),
        })
    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        return jsonify({"success": False, "error": str(e), "trace": tb[-1200:]})


@app.route("/security/debug_csv", methods=["POST"])
def security_debug_csv():
    """Debug: return raw column names + unique STATUS values from a Prowler CSV.
    POST body: {"csvPath": "<absolute path returned by /security/run>"}
    """
    b = request.json or {}
    csv_path = b.get("csvPath", "")
    if not csv_path or not os.path.exists(csv_path):
        return jsonify({"success": False, "error": f"CSV not found: {csv_path}"})
    try:
        import pandas as pd
        df = None
        used_sep = ","
        for sep in [",", ";", "\t", "|"]:
            for enc in ["utf-8", "utf-8-sig", "ISO-8859-1"]:
                try:
                    _df = pd.read_csv(csv_path, sep=sep, encoding=enc,
                                      nrows=500, on_bad_lines="skip")
                    if len(_df.columns) > 3:
                        df = _df; used_sep = sep; break
                except Exception:
                    continue
            if df is not None:
                break
        if df is None:
            return jsonify({"success": False, "error": "Could not parse CSV"})
        dup_cols = [c for c in df.columns if list(df.columns).count(c) > 1]
        status_cols = {}
        for c in df.columns:
            if "status" in c.lower():
                col_data = df[c]
                if isinstance(col_data, pd.DataFrame):
                    col_data = col_data.iloc[:, 0]
                status_cols[c] = col_data.dropna().unique().tolist()[:20]
        return jsonify({
            "success": True,
            "separator": repr(used_sep),
            "totalRows": len(df),
            "columns": list(df.columns),
            "duplicateColumns": list(set(dup_cols)),
            "statusColumns": status_cols,
            "sampleRow": df.head(1).fillna("").astype(str).to_dict("records"),
        })
    except Exception as e:
        import traceback
        return jsonify({"success": False, "error": str(e),
                        "trace": traceback.format_exc()[-500:]})


@app.route("/security/download_xlsx", methods=["POST"])
def security_download_xlsx():
    """Stream the server-generated XLSX file to the browser."""
    b = request.json or {}
    xlsx_path = b.get("xlsxPath", "")
    cust_name = b.get("customerName", "Report").replace(" ", "_")
    account_id = b.get("accountId", "acct")
    if not xlsx_path or not os.path.exists(xlsx_path):
        return jsonify({"success": False, "error": "XLSX file not found on server"}), 404
    fname = f"{cust_name}_{account_id}_security_audit.xlsx"
    return send_file(xlsx_path, as_attachment=True,
                     download_name=fname,
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# ─── MONTHLY REPORT ───────────────────────────────────────────
@app.route("/monthly/report", methods=["POST"])
def monthly_report():
    b = request.json or {}
    ak, sk = b.get("accessKey",""), b.get("secretKey","")
    regions_list = b.get("regions", ["us-east-1"])
    try:
        ce = boto3.client("ce", aws_access_key_id=ak, aws_secret_access_key=sk, region_name="us-east-1")
        now = datetime.utcnow()
        # Last 2 months
        m0_start = (now.replace(day=1) - timedelta(days=32)).replace(day=1)
        m1_start = now.replace(day=1)
        m0 = m0_start.strftime("%Y-%m")
        m1 = (now - timedelta(days=32)).replace(day=1).strftime("%Y-%m")
        m2 = now.replace(day=1).strftime("%Y-%m")
        start_str = m0_start.strftime("%Y-%m-%d")
        end_str = now.strftime("%Y-%m-%d")
        
        # Billing by service for last 2 months
        resp = ce.get_cost_and_usage(
            TimePeriod={"Start": start_str, "End": end_str},
            Granularity="MONTHLY",
            Metrics=["UnblendedCost"],
            GroupBy=[{"Type":"DIMENSION","Key":"SERVICE"}]
        )
        billing = {}
        months_found = []
        for r in resp["ResultsByTime"]:
            m = r["TimePeriod"]["Start"][:7]
            if m not in months_found: months_found.append(m)
            for grp in r["Groups"]:
                svc = grp["Keys"][0]
                cost = round(float(grp["Metrics"]["UnblendedCost"]["Amount"]), 4)
                if cost > 0:
                    if svc not in billing: billing[svc] = {}
                    billing[svc][m] = cost

        # Filter out tax/support noise
        EXCLUDE_M = ("tax", "support", "premium support", "refund", "credit",
                     "aws marketplace", "marketplace", "savings plan", "discount")
        billing = {k: v for k, v in billing.items()
                   if not any(p in k.lower() for p in EXCLUDE_M)}

        # Keep top 5 services by last-month cost
        last_m_billing = months_found[-1] if months_found else ""
        billing = dict(sorted(billing.items(),
            key=lambda x: x[1].get(last_m_billing, 0), reverse=True)[:5])
        
        # EC2 running instances with utilization (CloudWatch)
        ec2_util = []
        for region in regions_list:
            try:
                ec2 = boto3.client("ec2", aws_access_key_id=ak, aws_secret_access_key=sk, region_name=region)
                cw = boto3.client("cloudwatch", aws_access_key_id=ak, aws_secret_access_key=sk, region_name=region)
                paginator = ec2.get_paginator("describe_instances")
                cw_end = datetime.utcnow()
                cw_start = cw_end - timedelta(days=30)
                for page in paginator.paginate(Filters=[{"Name":"instance-state-name","Values":["running"]}]):
                    for res in page["Reservations"]:
                        for inst in res["Instances"]:
                            iid = inst["InstanceId"]
                            itype = inst.get("InstanceType","N/A")
                            name = next((t["Value"] for t in inst.get("Tags",[]) if t["Key"]=="Name"), "N/A")
                            # CPU
                            cpu_data = cw.get_metric_statistics(
                                Namespace="AWS/EC2", MetricName="CPUUtilization",
                                Dimensions=[{"Name":"InstanceId","Value":iid}],
                                StartTime=cw_start, EndTime=cw_end, Period=86400*30,
                                Statistics=["Average","Maximum"]
                            )
                            cpu_avg = round(cpu_data["Datapoints"][0]["Average"],1) if cpu_data["Datapoints"] else "N/A"
                            cpu_max = round(cpu_data["Datapoints"][0]["Maximum"],1) if cpu_data["Datapoints"] else "N/A"
                            ec2_util.append({"Region":region,"Name":name,"Instance ID":iid,
                                "Type":itype,"CPU Avg (%)":cpu_avg,"CPU Max (%)":cpu_max,
                                "Memory Avg (%)":"N/A*","Disk Read (MB)":"N/A*"})
            except: pass
        
        # RDS instances
        rds_util = []
        for region in regions_list:
            try:
                rds = boto3.client("rds", aws_access_key_id=ak, aws_secret_access_key=sk, region_name=region)
                cw = boto3.client("cloudwatch", aws_access_key_id=ak, aws_secret_access_key=sk, region_name=region)
                dbs = rds.describe_db_instances()["DBInstances"]
                cw_end = datetime.utcnow()
                cw_start = cw_end - timedelta(days=30)
                for db in dbs:
                    dbid = db["DBInstanceIdentifier"]
                    cpu_data = cw.get_metric_statistics(
                        Namespace="AWS/RDS", MetricName="CPUUtilization",
                        Dimensions=[{"Name":"DBInstanceIdentifier","Value":dbid}],
                        StartTime=cw_start, EndTime=cw_end, Period=86400*30,
                        Statistics=["Average","Maximum"]
                    )
                    cpu_avg = round(cpu_data["Datapoints"][0]["Average"],1) if cpu_data["Datapoints"] else "N/A"
                    cpu_max = round(cpu_data["Datapoints"][0]["Maximum"],1) if cpu_data["Datapoints"] else "N/A"
                    rds_util.append({"Region":region,"DB ID":dbid,"Engine":db.get("Engine","N/A"),
                        "Class":db.get("DBInstanceClass","N/A"),"Status":db.get("DBInstanceStatus","N/A"),
                        "CPU Avg (%)":cpu_avg,"CPU Max (%)":cpu_max,"Storage (GB)":db.get("AllocatedStorage","N/A")})
            except: pass
        
        return jsonify({"success": True, "billingMonths": months_found, "billing": billing,
            "ec2Utilization": ec2_util, "rdsUtilization": rds_util})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/monthly/export_docx", methods=["POST"])
def monthly_export_docx():
    """Generate a .docx monthly report matching the reference format."""
    b = request.json or {}
    cust     = b.get("customerName", "Report")
    acct     = b.get("accountId", "")
    acc_name = b.get("accName", acct)
    data     = b.get("data", {})

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm, Inches, Twips
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import io

        doc = Document()

        # ── Page margins ──────────────────────────────────────
        for section in doc.sections:
            section.top_margin    = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin   = Cm(1.8)
            section.right_margin  = Cm(1.8)

        # ── Helpers ───────────────────────────────────────────
        def set_cell_bg(cell, hex_colour):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), hex_colour.lstrip("#")); tcPr.append(shd)

        def set_cell_border(cell):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for side in ("top","left","bottom","right"):
                b_el = OxmlElement(f"w:{side}")
                b_el.set(qn("w:val"),"single"); b_el.set(qn("w:sz"),"4")
                b_el.set(qn("w:color"),"BFCFE8"); borders.append(b_el)
            tcPr.append(borders)

        def para(text, bold=False, size=10, colour="1E293B",
                 align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
            p = doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after  = Pt(space_after)
            r = p.add_run(text); r.bold = bold
            r.font.size = Pt(size); r.font.color.rgb = RGBColor.from_string(colour)
            return p

        def section_heading(num, title, colour="1A3A6B"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(f"{num}. {title}")
            r.bold = True; r.font.size = Pt(12)
            r.font.color.rgb = RGBColor.from_string(colour)
            # Bottom border via pPr
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            btm = OxmlElement("w:bottom")
            btm.set(qn("w:val"),"single"); btm.set(qn("w:sz"),"6")
            btm.set(qn("w:color"),"1A3A6B"); pBdr.append(btm); pPr.append(pBdr)
            return p

        def styled_table(headers, rows, hdr_bg="1A3A6B", hdr_fg="FFFFFF",
                         col_widths_cm=None, right_align_from=1):
            n = len(headers)
            tbl = doc.add_table(rows=1 + len(rows), cols=n)
            tbl.style = "Table Grid"; tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            # Header
            hr = tbl.rows[0]
            for ci, h in enumerate(headers):
                cell = hr.cells[ci]; set_cell_bg(cell, hdr_bg)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci >= right_align_from else WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(str(h)); r.bold = True
                r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string(hdr_fg)
            # Data
            for ri, row_data in enumerate(rows):
                tr = tbl.rows[ri + 1]
                is_total = str(row_data.get(headers[0],"")).upper().startswith("TOTAL")
                bg = "E8EFF8" if is_total else ("F5F8FD" if ri % 2 == 0 else "FFFFFF")
                for ci, h in enumerate(headers):
                    cell = tr.cells[ci]; set_cell_bg(cell, bg); set_cell_border(cell)
                    val = str(row_data.get(h, ""))
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if ci >= right_align_from else WD_ALIGN_PARAGRAPH.LEFT
                    r = p.add_run(val); r.font.size = Pt(8.5)
                    r.bold = is_total
                    r.font.color.rgb = RGBColor.from_string("1A3A6B" if is_total else "1E293B")
            # Column widths
            if col_widths_cm:
                for ci, w in enumerate(col_widths_cm):
                    for row in tbl.rows:
                        row.cells[ci].width = Cm(w)
            return tbl

        # ── Data extraction ────────────────────────────────────
        now_str  = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        now_disp = datetime.now().strftime("%B %Y")
        months   = data.get("billingMonths", [])
        m0 = months[0] if len(months) > 0 else "Month-1"
        m1 = months[1] if len(months) > 1 else "Month-2"
        billing  = data.get("billing", {})
        ec2_util = data.get("ec2Utilization", [])
        rds_util = data.get("rdsUtilization", [])

        # ════════════════════════════════════════════════════════
        # COVER / TITLE BLOCK
        # ════════════════════════════════════════════════════════
        # Dark banner paragraph (simulate with a 1-row, 1-col table)
        banner_tbl = doc.add_table(rows=1, cols=1)
        banner_tbl.style = "Table Grid"
        banner_cell = banner_tbl.rows[0].cells[0]
        set_cell_bg(banner_cell, "0C1929")
        bp = banner_cell.paragraphs[0]
        bp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        bp.paragraph_format.space_before = Pt(8); bp.paragraph_format.space_after = Pt(2)
        r1 = bp.add_run(f"{cust} Monthly Report  ·  {now_disp}")
        r1.bold = True; r1.font.size = Pt(16); r1.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        # second line
        bp2 = banner_cell.add_paragraph()
        bp2.paragraph_format.space_before = Pt(2); bp2.paragraph_format.space_after = Pt(8)
        r2 = bp2.add_run(f"Account ID: {acct}   |   Account Name: {acc_name}   |   Generated: {now_str}")
        r2.font.size = Pt(8.5); r2.font.color.rgb = RGBColor(0xAE,0xC9,0xF0)

        doc.add_paragraph()  # spacer

        # Compute totals for cover stats
        total_m0 = sum(v.get(m0,0) for v in billing.values())
        total_m1 = sum(v.get(m1,0) for v in billing.values())
        total_diff = total_m1 - total_m0

        # Quick-stats row (2-col mini table)
        stats_tbl = doc.add_table(rows=1, cols=3)
        stats_tbl.style = "Table Grid"; stats_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        stat_data = [
            (f"${total_m0:,.2f}", m0 + " Total"),
            (f"${total_m1:,.2f}", m1 + " Total"),
            (f"{total_diff:+,.2f}", "Difference ($)"),
        ]
        for ci, (val, label) in enumerate(stat_data):
            cell = stats_tbl.rows[0].cells[ci]
            set_cell_bg(cell, "1A3A6B")
            cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before = Pt(6); cp.paragraph_format.space_after = Pt(2)
            rv = cp.add_run(val); rv.bold = True
            rv.font.size = Pt(14); rv.font.color.rgb = RGBColor(0xF5,0x9E,0x0B)
            lp = cell.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lp.paragraph_format.space_before = Pt(0); lp.paragraph_format.space_after = Pt(6)
            rl = lp.add_run(label); rl.font.size = Pt(8)
            rl.font.color.rgb = RGBColor(0xAE,0xC9,0xF0)

        doc.add_paragraph()

        # ════════════════════════════════════════════════════════
        # SECTION 1 — BILLING COMPARISON (TOP 5)
        # ════════════════════════════════════════════════════════
        section_heading(1, f"Billing Comparison (Top 5 Services) — {m0} vs {m1}")

        billing_rows = []
        for svc, monthly in sorted(billing.items(),
                key=lambda x: x[1].get(m1, 0), reverse=True):
            c0 = monthly.get(m0, 0) or 0
            c1 = monthly.get(m1, 0) or 0
            if c0 < 0.01 and c1 < 0.01: continue
            diff = c1 - c0
            pct  = f"{((diff/c0)*100):+.1f}%" if c0 > 0 else "—"
            billing_rows.append({
                "Service": svc,
                f"{m0} ($)": f"${c0:,.2f}",
                f"{m1} ($)": f"${c1:,.2f}",
                "Difference ($)": f"${diff:+,.2f}",
                "Change (%)": pct,
            })
        billing_rows.append({
            "Service": "TOTAL",
            f"{m0} ($)": f"${total_m0:,.2f}",
            f"{m1} ($)": f"${total_m1:,.2f}",
            "Difference ($)": f"${total_diff:+,.2f}",
            "Change (%)": f"{((total_diff/total_m0)*100):+.1f}%" if total_m0 > 0 else "—",
        })
        billing_hdrs = ["Service", f"{m0} ($)", f"{m1} ($)", "Difference ($)", "Change (%)"]
        styled_table(billing_hdrs, billing_rows,
                     col_widths_cm=[7.5, 3.0, 3.0, 3.0, 2.5])
        doc.add_paragraph()

        # ════════════════════════════════════════════════════════
        # SECTION 2 — EC2 INSTANCE UTILIZATION
        # ════════════════════════════════════════════════════════
        section_heading(2, "EC2 Instance Utilization — Last 30 Days", colour="243C72")
        para("CPU utilization from CloudWatch. Memory/Disk require CloudWatch Agent on instances.",
             size=8.5, colour="64748B", space_after=6)
        if ec2_util:
            styled_table(list(ec2_util[0].keys()), ec2_util, hdr_bg="243C72",
                         col_widths_cm=None, right_align_from=4)
        else:
            para("No running EC2 instances found in the selected regions.", colour="94A3B8")
        doc.add_paragraph()

        # ════════════════════════════════════════════════════════
        # SECTION 3 — RDS UTILIZATION
        # ════════════════════════════════════════════════════════
        section_heading(3, "RDS Instance Utilization — Last 30 Days", colour="1C3560")
        if rds_util:
            styled_table(list(rds_util[0].keys()), rds_util, hdr_bg="1C3560",
                         col_widths_cm=None, right_align_from=4)
        else:
            para("No RDS instances found in the selected regions.", colour="94A3B8")
        doc.add_paragraph()

        # ════════════════════════════════════════════════════════
        # FOOTER
        # ════════════════════════════════════════════════════════
        doc.add_paragraph()
        fp = doc.add_paragraph()
        fp.paragraph_format.space_before = Pt(10)
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run(f"Generated by AWS Cloud Suite  ·  Customer: {cust}  ·  Account: {acct}  ·  {now_str}")
        fr.font.size = Pt(7.5); fr.font.color.rgb = RGBColor(0x94,0xA3,0xB8)

        # ── Stream response ────────────────────────────────────
        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        filename = f"{cust.replace(' ','_')}_{acct}_monthly_report.docx"
        return send_file(buf,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True, download_name=filename)

    except ImportError:
        return jsonify({"success": False,
            "error": "python-docx not installed. Run: pip install python-docx"}), 500
    except Exception as e:
        import traceback
        return jsonify({"success": False, "error": str(e),
            "trace": traceback.format_exc()[-600:]}), 500


# ─── Helpers ──────────────────────────────────────────────────
def sess(ak, sk, region="us-east-1"):
    return boto3.Session(aws_access_key_id=ak, aws_secret_access_key=sk, region_name=region)

def cli(s, svc, region=None):
    return s.client(svc, region_name=region or s.region_name)

def safe(v):
    if v is None: return "N/A"
    if isinstance(v, bool): return "Yes" if v else "No"
    if isinstance(v, (int, float)): return v
    if isinstance(v, list): return ", ".join(str(i) for i in v) if v else "N/A"
    if isinstance(v, dict): return json.dumps(v)
    return str(v)

def tag_name(tags):
    if not tags: return "N/A"
    for t in tags:
        if t.get("Key") == "Name": return t.get("Value", "N/A")
    return "N/A"


# ════════════════════════════════════════════════
#   COLLECTORS — 200+ AWS Services
# ════════════════════════════════════════════════

def collect_ec2(s, region):
    ec2 = cli(s,"ec2",region); rows = []
    for page in ec2.get_paginator("describe_instances").paginate():
        for res in page["Reservations"]:
            for i in res["Instances"]:
                vols = i.get("BlockDeviceMappings", [])
                vol_ids = [v["Ebs"]["VolumeId"] for v in vols if "Ebs" in v]
                vol_size = "N/A"
                if vol_ids:
                    try:
                        vd = ec2.describe_volumes(VolumeIds=vol_ids[:1])["Volumes"]
                        vol_size = str(vd[0]["Size"]) + " GB" if vd else "N/A"
                    except: pass
                sgs = i.get("SecurityGroups", [])
                rows.append({
                    "Instance Name": tag_name(i.get("Tags")),
                    "Instance ID": safe(i.get("InstanceId")),
                    "Instance Type": safe(i.get("InstanceType")),
                    "State": safe(i.get("State",{}).get("Name")),
                    "AMI ID": safe(i.get("ImageId")),
                    "Key Pair": safe(i.get("KeyName")),
                    "VPC ID": safe(i.get("VpcId")),
                    "Subnet ID": safe(i.get("SubnetId")),
                    "Availability Zone": safe(i.get("Placement",{}).get("AvailabilityZone")),
                    "Private IP": safe(i.get("PrivateIpAddress")),
                    "Public IP": safe(i.get("PublicIpAddress")),
                    "Volume IDs": safe(vol_ids),
                    "Root Volume Size": vol_size,
                    "Volume Count": len(vols),
                    "Security Group IDs": safe([sg["GroupId"] for sg in sgs]),
                    "Security Group Names": safe([sg["GroupName"] for sg in sgs]),
                    "IAM Role": safe(i.get("IamInstanceProfile",{}).get("Arn","N/A").split("/")[-1] if i.get("IamInstanceProfile") else "N/A"),
                    "Platform": safe(i.get("Platform","Linux/UNIX")),
                    "Architecture": safe(i.get("Architecture")),
                    "Launch Time": safe(str(i.get("LaunchTime",""))[:19]),
                    "EBS Optimized": safe(i.get("EbsOptimized")),
                    "Monitoring": safe(i.get("Monitoring",{}).get("State")),
                })
    return rows

def collect_ec2_ami(s, region):
    ec2 = cli(s,"ec2",region)
    images = ec2.describe_images(Owners=["self"])["Images"]
    return [{"Image Name":safe(i.get("Name")),"Image ID":safe(i.get("ImageId")),
             "State":safe(i.get("State")),"Architecture":safe(i.get("Architecture")),
             "Virtualization":safe(i.get("VirtualizationType")),"Public":safe(i.get("Public")),
             "Owner ID":safe(i.get("OwnerId")),"Platform":safe(i.get("PlatformDetails")),
             "Creation Date":safe(i.get("CreationDate","")[:19]),"Description":safe(i.get("Description","N/A"))} for i in images]

def collect_ec2_sg(s, region):
    ec2 = cli(s,"ec2",region); rows = []
    for page in ec2.get_paginator("describe_security_groups").paginate():
        for sg in page["SecurityGroups"]:
            rows.append({"Group Name":safe(sg.get("GroupName")),"Group ID":safe(sg.get("GroupId")),
                         "Description":safe(sg.get("Description")),"VPC ID":safe(sg.get("VpcId")),
                         "Owner ID":safe(sg.get("OwnerId")),"Inbound Rules":len(sg.get("IpPermissions",[])),
                         "Outbound Rules":len(sg.get("IpPermissionsEgress",[])),"Tags":tag_name(sg.get("Tags"))})
    return rows

def collect_ec2_keypair(s, region):
    ec2 = cli(s,"ec2",region)
    pairs = ec2.describe_key_pairs()["KeyPairs"]
    return [{"Key Name":safe(kp.get("KeyName")),"Key ID":safe(kp.get("KeyPairId")),
             "Fingerprint":safe(kp.get("KeyFingerprint")),"Type":safe(kp.get("KeyType","rsa")),
             "Creation Date":safe(str(kp.get("CreateTime",""))[:19]),"Tags":tag_name(kp.get("Tags"))} for kp in pairs]

def collect_ec2_eip(s, region):
    ec2 = cli(s,"ec2",region)
    addrs = ec2.describe_addresses()["Addresses"]
    return [{"Public IP":safe(a.get("PublicIp")),"Allocation ID":safe(a.get("AllocationId")),
             "Domain":safe(a.get("Domain")),"Instance ID":safe(a.get("InstanceId","N/A")),
             "Association ID":safe(a.get("AssociationId","N/A")),"Private IP":safe(a.get("PrivateIpAddress","N/A")),
             "Tags":tag_name(a.get("Tags"))} for a in addrs]

def collect_ec2_snapshot(s, region):
    ec2 = cli(s,"ec2",region); rows = []
    for page in ec2.get_paginator("describe_snapshots").paginate(OwnerIds=["self"]):
        for snap in page["Snapshots"]:
            rows.append({"Snapshot ID":safe(snap.get("SnapshotId")),"Volume ID":safe(snap.get("VolumeId")),
                         "State":safe(snap.get("State")),"Volume Size (GB)":safe(snap.get("VolumeSize")),
                         "Encrypted":safe(snap.get("Encrypted")),"Owner ID":safe(snap.get("OwnerId")),
                         "Progress":safe(snap.get("Progress")),"Start Time":safe(str(snap.get("StartTime",""))[:19]),
                         "Tags":tag_name(snap.get("Tags"))})
    return rows

def collect_ec2_volume(s, region):
    ec2 = cli(s,"ec2",region); rows = []
    for page in ec2.get_paginator("describe_volumes").paginate():
        for v in page["Volumes"]:
            att = v.get("Attachments", [])
            rows.append({"Volume ID":safe(v.get("VolumeId")),"Volume Name":tag_name(v.get("Tags")),
                         "State":safe(v.get("State")),"Size (GB)":safe(v.get("Size")),
                         "Volume Type":safe(v.get("VolumeType")),"IOPS":safe(v.get("Iops","N/A")),
                         "Encrypted":safe(v.get("Encrypted")),"AZ":safe(v.get("AvailabilityZone")),
                         "Attached Instance":safe(att[0].get("InstanceId") if att else "N/A"),
                         "Attach State":safe(att[0].get("State") if att else "N/A"),
                         "Creation Time":safe(str(v.get("CreateTime",""))[:19])})
    return rows

def collect_autoscaling(s, region):
    asg = cli(s,"autoscaling",region); rows = []
    for page in asg.get_paginator("describe_auto_scaling_groups").paginate():
        for g in page["AutoScalingGroups"]:
            rows.append({"Group Name":safe(g.get("AutoScalingGroupName")),"ARN":safe(g.get("AutoScalingGroupARN")),
                         "Min Size":safe(g.get("MinSize")),"Max Size":safe(g.get("MaxSize")),
                         "Desired":safe(g.get("DesiredCapacity")),"Current Instances":len(g.get("Instances",[])),
                         "Health Check Type":safe(g.get("HealthCheckType")),"AZs":safe(g.get("AvailabilityZones",[])),
                         "Created":safe(str(g.get("CreatedTime",""))[:19])})
    return rows

def collect_lightsail(s, region):
    try:
        ls = cli(s,"lightsail",region)
        instances = ls.get_instances()["instances"]
        return [{"Name":safe(i.get("name")),"State":safe(i.get("state",{}).get("name")),
                 "Blueprint":safe(i.get("blueprintId")),"Bundle":safe(i.get("bundleId")),
                 "Public IP":safe(i.get("publicIpAddress","N/A")),"Private IP":safe(i.get("privateIpAddress","N/A")),
                 "Created":safe(str(i.get("createdAt",""))[:19])} for i in instances]
    except: return []

def collect_batch(s, region):
    try:
        b = cli(s,"batch",region)
        envs = b.describe_compute_environments()["computeEnvironments"]
        return [{"Env Name":safe(e.get("computeEnvironmentName")),"ARN":safe(e.get("computeEnvironmentArn")),
                 "Type":safe(e.get("type")),"State":safe(e.get("state")),"Status":safe(e.get("status"))} for e in envs]
    except: return []

def collect_s3(s, region):
    s3 = cli(s,"s3"); rows = []
    for b in s3.list_buckets().get("Buckets",[]):
        name = b["Name"]
        try: bregion = s3.get_bucket_location(Bucket=name).get("LocationConstraint") or "us-east-1"
        except: bregion = "N/A"
        try: ver = s3.get_bucket_versioning(Bucket=name).get("Status","Disabled") or "Disabled"
        except: ver = "N/A"
        try:
            enc = s3.get_bucket_encryption(Bucket=name)
            sse = enc["ServerSideEncryptionConfiguration"]["Rules"][0]["ApplyServerSideEncryptionByDefault"]["SSEAlgorithm"]
        except: sse = "Not Enabled"
        try:
            pub = s3.get_public_access_block(Bucket=name)["PublicAccessBlockConfiguration"]
            block = "All Blocked" if all([pub.get("BlockPublicAcls"),pub.get("IgnorePublicAcls"),pub.get("BlockPublicPolicy"),pub.get("RestrictPublicBuckets")]) else "Partially Open"
        except: block = "N/A"
        rows.append({"Bucket Name":name,"Region":bregion,"Creation Date":safe(str(b.get("CreationDate",""))[:19]),
                     "Versioning":ver,"Encryption":sse,"Public Access Block":block})
    return rows

def collect_efs(s, region):
    try:
        efs = cli(s,"efs",region)
        fss = efs.describe_file_systems()["FileSystems"]
        return [{"File System ID":safe(f.get("FileSystemId")),"Name":tag_name(f.get("Tags")),
                 "Lifecycle State":safe(f.get("LifeCycleState")),"Performance Mode":safe(f.get("PerformanceMode")),
                 "Encrypted":safe(f.get("Encrypted")),"Size (bytes)":safe(f.get("SizeInBytes",{}).get("Value")),
                 "Mount Targets":safe(f.get("NumberOfMountTargets")),"ARN":safe(f.get("FileSystemArn")),
                 "Creation Time":safe(str(f.get("CreationTime",""))[:19])} for f in fss]
    except: return []

def collect_fsx(s, region):
    try:
        fsx = cli(s,"fsx",region)
        fss = fsx.describe_file_systems()["FileSystems"]
        return [{"File System ID":safe(f.get("FileSystemId")),"Type":safe(f.get("FileSystemType")),
                 "Lifecycle":safe(f.get("Lifecycle")),"Storage (GB)":safe(f.get("StorageCapacity")),
                 "VPC ID":safe(f.get("VpcId")),"DNS Name":safe(f.get("DNSName","N/A")),
                 "Created":safe(str(f.get("CreationTime",""))[:19])} for f in fss]
    except: return []

def collect_backup(s, region):
    try:
        b = cli(s,"backup",region)
        vaults = b.list_backup_vaults()["BackupVaultList"]
        return [{"Vault Name":safe(v.get("BackupVaultName")),"ARN":safe(v.get("BackupVaultArn")),
                 "Recovery Points":safe(v.get("NumberOfRecoveryPoints")),
                 "Created":safe(str(v.get("CreationDate",""))[:19])} for v in vaults]
    except: return []

def collect_rds(s, region):
    rds = cli(s,"rds",region); rows = []
    for page in rds.get_paginator("describe_db_instances").paginate():
        for db in page["DBInstances"]:
            rows.append({
                "DB Instance ID":safe(db.get("DBInstanceIdentifier")),"Engine":safe(db.get("Engine")),
                "Engine Version":safe(db.get("EngineVersion")),"Instance Class":safe(db.get("DBInstanceClass")),
                "Status":safe(db.get("DBInstanceStatus")),"Multi-AZ":safe(db.get("MultiAZ")),
                "AZ":safe(db.get("AvailabilityZone")),"Storage Type":safe(db.get("StorageType")),
                "Storage (GB)":safe(db.get("AllocatedStorage")),"VPC ID":safe(db.get("DBSubnetGroup",{}).get("VpcId")),
                "Endpoint":safe(db.get("Endpoint",{}).get("Address")),"Port":safe(db.get("Endpoint",{}).get("Port")),
                "DB Name":safe(db.get("DBName")),"Master User":safe(db.get("MasterUsername")),
                "Backup Retention":safe(db.get("BackupRetentionPeriod")),"Encrypted":safe(db.get("StorageEncrypted")),
                "Deletion Protection":safe(db.get("DeletionProtection")),"IAM Auth":safe(db.get("IAMDatabaseAuthenticationEnabled")),
                "ARN":safe(db.get("DBInstanceArn")),"Created":safe(str(db.get("InstanceCreateTime",""))[:19])
            })
    return rows

def collect_rds_snapshot(s, region):
    rds = cli(s,"rds",region); rows = []
    for page in rds.get_paginator("describe_db_snapshots").paginate(SnapshotType="manual"):
        for snap in page["DBSnapshots"]:
            rows.append({"Snapshot ID":safe(snap.get("DBSnapshotIdentifier")),"DB Instance":safe(snap.get("DBInstanceIdentifier")),
                         "Status":safe(snap.get("Status")),"Engine":safe(snap.get("Engine")),
                         "Storage (GB)":safe(snap.get("AllocatedStorage")),"Encrypted":safe(snap.get("Encrypted")),
                         "Created":safe(str(snap.get("SnapshotCreateTime",""))[:19])})
    return rows

def collect_dynamodb(s, region):
    ddb = cli(s,"dynamodb",region); rows = []
    for page in ddb.get_paginator("list_tables").paginate():
        for tname in page["TableNames"]:
            try:
                t = ddb.describe_table(TableName=tname)["Table"]
                billing = t.get("BillingModeSummary",{}).get("BillingMode","PROVISIONED")
                tp = t.get("ProvisionedThroughput",{})
                rows.append({"Table Name":tname,"Status":safe(t.get("TableStatus")),
                             "Billing Mode":billing,"Item Count":safe(t.get("ItemCount")),
                             "Size (bytes)":safe(t.get("TableSizeBytes")),
                             "Partition Key":safe(t["KeySchema"][0]["AttributeName"]) if t.get("KeySchema") else "N/A",
                             "Sort Key":safe(t["KeySchema"][1]["AttributeName"]) if len(t.get("KeySchema",[]))>1 else "N/A",
                             "Read Capacity":safe(tp.get("ReadCapacityUnits","N/A")),
                             "Write Capacity":safe(tp.get("WriteCapacityUnits","N/A")),
                             "Stream Enabled":safe(t.get("StreamSpecification",{}).get("StreamEnabled","No")),
                             "ARN":safe(t.get("TableArn")),"Created":safe(str(t.get("CreationDateTime",""))[:19])})
            except: rows.append({"Table Name":tname,"Error":"Failed to describe"})
    return rows

def collect_elasticache(s, region):
    ec = cli(s,"elasticache",region); rows = []
    for page in ec.get_paginator("describe_cache_clusters").paginate(ShowCacheNodeInfo=True):
        for c in page["CacheClusters"]:
            ep = c.get("ConfigurationEndpoint") or (c.get("CacheNodes",[{}])[0].get("Endpoint",{}) if c.get("CacheNodes") else {})
            rows.append({"Cluster ID":safe(c.get("CacheClusterId")),"Engine":safe(c.get("Engine")),
                         "Engine Version":safe(c.get("EngineVersion")),"Node Type":safe(c.get("CacheNodeType")),
                         "Status":safe(c.get("CacheClusterStatus")),"Num Nodes":safe(c.get("NumCacheNodes")),
                         "Endpoint":safe(ep.get("Address","N/A")),"Port":safe(ep.get("Port","N/A")),
                         "Subnet Group":safe(c.get("CacheSubnetGroupName","N/A")),
                         "Created":safe(str(c.get("CacheClusterCreateTime",""))[:19])})
    return rows

def collect_redshift(s, region):
    try:
        rs = cli(s,"redshift",region)
        clusters = rs.describe_clusters()["Clusters"]
        return [{"Cluster ID":safe(c.get("ClusterIdentifier")),"Status":safe(c.get("ClusterStatus")),
                 "Node Type":safe(c.get("NodeType")),"Nodes":safe(c.get("NumberOfNodes")),
                 "DB Name":safe(c.get("DBName")),"Endpoint":safe(c.get("Endpoint",{}).get("Address","N/A")),
                 "VPC ID":safe(c.get("VpcId","N/A")),"Encrypted":safe(c.get("Encrypted")),
                 "Created":safe(str(c.get("ClusterCreateTime",""))[:19])} for c in clusters]
    except: return []

def collect_docdb(s, region):
    try:
        rds = cli(s,"docdb",region)
        clusters = rds.describe_db_clusters()["DBClusters"]
        return [{"Cluster ID":safe(c.get("DBClusterIdentifier")),"Status":safe(c.get("Status")),
                 "Engine":safe(c.get("Engine")),"Engine Version":safe(c.get("EngineVersion")),
                 "Endpoint":safe(c.get("Endpoint","N/A")),"Port":safe(c.get("Port")),
                 "Multi-AZ":safe(c.get("MultiAZ")),"Encrypted":safe(c.get("StorageEncrypted")),
                 "ARN":safe(c.get("DBClusterArn")),"Created":safe(str(c.get("ClusterCreateTime",""))[:19])} for c in clusters]
    except: return []

def collect_neptune(s, region):
    try:
        rds = cli(s,"neptune",region)
        clusters = rds.describe_db_clusters()["DBClusters"]
        return [{"Cluster ID":safe(c.get("DBClusterIdentifier")),"Status":safe(c.get("Status")),
                 "Engine Version":safe(c.get("EngineVersion")),"Endpoint":safe(c.get("Endpoint","N/A")),
                 "Multi-AZ":safe(c.get("MultiAZ")),"Encrypted":safe(c.get("StorageEncrypted")),
                 "ARN":safe(c.get("DBClusterArn"))} for c in clusters if c.get("Engine","").startswith("neptune")]
    except: return []

def collect_vpc(s, region):
    ec2 = cli(s,"ec2",region); rows = []
    vpcs = ec2.describe_vpcs()["Vpcs"]
    for vpc in vpcs:
        vid = vpc["VpcId"]
        try:
            dns_host = safe(ec2.describe_vpc_attribute(VpcId=vid,Attribute="enableDnsHostnames")["EnableDnsHostnames"]["Value"])
        except: dns_host = "N/A"
        rows.append({"VPC ID":vid,"VPC Name":tag_name(vpc.get("Tags")),"CIDR Block":safe(vpc.get("CidrBlock")),
                     "State":safe(vpc.get("State")),"Is Default":safe(vpc.get("IsDefault")),
                     "DNS Hostnames":dns_host,"Owner ID":safe(vpc.get("OwnerId"))})
    return rows

def collect_subnet(s, region):
    ec2 = cli(s,"ec2",region); rows = []
    for page in ec2.get_paginator("describe_subnets").paginate():
        for sub in page["Subnets"]:
            rows.append({"Subnet ID":safe(sub.get("SubnetId")),"Name":tag_name(sub.get("Tags")),
                         "VPC ID":safe(sub.get("VpcId")),"CIDR":safe(sub.get("CidrBlock")),
                         "State":safe(sub.get("State")),"AZ":safe(sub.get("AvailabilityZone")),
                         "Available IPs":safe(sub.get("AvailableIpAddressCount")),
                         "Auto-assign Public IP":safe(sub.get("MapPublicIpOnLaunch")),
                         "ARN":safe(sub.get("SubnetArn"))})
    return rows

def collect_routetable(s, region):
    ec2 = cli(s,"ec2",region); rows = []
    for page in ec2.get_paginator("describe_route_tables").paginate():
        for rt in page["RouteTables"]:
            assoc = rt.get("Associations",[])
            rows.append({"Route Table ID":safe(rt.get("RouteTableId")),"Name":tag_name(rt.get("Tags")),
                         "VPC ID":safe(rt.get("VpcId")),"Is Main":safe(any(a.get("Main") for a in assoc)),
                         "Routes Count":len(rt.get("Routes",[])),"Owner":safe(rt.get("OwnerId"))})
    return rows

def collect_igw(s, region):
    ec2 = cli(s,"ec2",region)
    igws = ec2.describe_internet_gateways()["InternetGateways"]
    return [{"IGW ID":safe(ig.get("InternetGatewayId")),"Name":tag_name(ig.get("Tags")),
             "State":safe(ig.get("Attachments",[{}])[0].get("State","detached") if ig.get("Attachments") else "detached"),
             "VPC ID":safe(ig.get("Attachments",[{}])[0].get("VpcId","N/A") if ig.get("Attachments") else "N/A"),
             "Owner":safe(ig.get("OwnerId"))} for ig in igws]

def collect_nat(s, region):
    ec2 = cli(s,"ec2",region); rows = []
    for page in ec2.get_paginator("describe_nat_gateways").paginate():
        for n in page["NatGateways"]:
            addr = n.get("NatGatewayAddresses",[{}])[0] if n.get("NatGatewayAddresses") else {}
            rows.append({"NAT Gateway ID":safe(n.get("NatGatewayId")),"Name":tag_name(n.get("Tags")),
                         "State":safe(n.get("State")),"Type":safe(n.get("ConnectivityType")),
                         "VPC ID":safe(n.get("VpcId")),"Subnet ID":safe(n.get("SubnetId")),
                         "Public IP":safe(addr.get("PublicIp","N/A")),"Private IP":safe(addr.get("PrivateIp","N/A")),
                         "Created":safe(str(n.get("CreateTime",""))[:19])})
    return rows

def collect_vpn(s, region):
    try:
        ec2 = cli(s,"ec2",region)
        vpns = ec2.describe_vpn_connections()["VpnConnections"]
        return [{"VPN ID":safe(v.get("VpnConnectionId")),"Name":tag_name(v.get("Tags")),
                 "State":safe(v.get("State")),"Type":safe(v.get("Type")),
                 "Customer GW":safe(v.get("CustomerGatewayId")),"VPN GW":safe(v.get("VpnGatewayId","N/A"))} for v in vpns]
    except: return []

def collect_directconnect(s, region):
    try:
        dc = cli(s,"directconnect",region)
        conns = dc.describe_connections()["connections"]
        return [{"Connection ID":safe(c.get("connectionId")),"Name":safe(c.get("connectionName")),
                 "State":safe(c.get("connectionState")),"Bandwidth":safe(c.get("bandwidth")),
                 "Location":safe(c.get("location")),"Owner":safe(c.get("ownerAccount"))} for c in conns]
    except: return []

def collect_elb(s, region):
    rows = []
    try:
        elbv2 = cli(s,"elbv2",region)
        for page in elbv2.get_paginator("describe_load_balancers").paginate():
            for lb in page["LoadBalancers"]:
                rows.append({"Name":safe(lb.get("LoadBalancerName")),"ARN":safe(lb.get("LoadBalancerArn")),
                             "Type":safe(lb.get("Type")),"Scheme":safe(lb.get("Scheme")),
                             "State":safe(lb.get("State",{}).get("Code")),
                             "DNS Name":safe(lb.get("DNSName")),"VPC ID":safe(lb.get("VpcId")),
                             "AZs":safe([az["ZoneName"] for az in lb.get("AvailabilityZones",[])]),"IP Type":safe(lb.get("IpAddressType")),
                             "Created":safe(str(lb.get("CreatedTime",""))[:19])})
    except Exception as e:
        rows.append({"Error": str(e)})
    return rows

def collect_cloudfront(s, region):
    cf = cli(s,"cloudfront","us-east-1"); rows = []
    for page in cf.get_paginator("list_distributions").paginate():
        for d in page.get("DistributionList",{}).get("Items",[]):
            rows.append({"Distribution ID":safe(d.get("Id")),"Domain":safe(d.get("DomainName")),
                         "Status":safe(d.get("Status")),"State":"Enabled" if d.get("Enabled") else "Disabled",
                         "Price Class":safe(d.get("PriceClass")),"Comment":safe(d.get("Comment","N/A")),
                         "Last Modified":safe(str(d.get("LastModifiedTime",""))[:19])})
    return rows

def collect_route53(s, region):
    r53 = cli(s,"route53","us-east-1"); rows = []
    for page in r53.get_paginator("list_hosted_zones").paginate():
        for z in page["HostedZones"]:
            zid = z["Id"].split("/")[-1]
            try: record_count = len(r53.list_resource_record_sets(HostedZoneId=zid).get("ResourceRecordSets",[]))
            except: record_count = "N/A"
            rows.append({"Zone Name":safe(z.get("Name")),"Zone ID":zid,
                         "Type":"Private" if z.get("Config",{}).get("PrivateZone") else "Public",
                         "Record Count":record_count,"Comment":safe(z.get("Config",{}).get("Comment","N/A"))})
    return rows

def collect_route53_record(s, region):
    r53 = cli(s,"route53","us-east-1"); rows = []
    try:
        for page in r53.get_paginator("list_hosted_zones").paginate():
            for z in page["HostedZones"]:
                zid = z["Id"].split("/")[-1]; zname = z["Name"]
                try:
                    for rr in r53.list_resource_record_sets(HostedZoneId=zid).get("ResourceRecordSets",[]):
                        rows.append({"Zone":zname,"Record Name":safe(rr.get("Name")),
                                     "Type":safe(rr.get("Type")),"TTL":safe(rr.get("TTL","N/A")),
                                     "Alias":"Yes" if rr.get("AliasTarget") else "No",
                                     "Values":safe([r.get("Value","") for r in rr.get("ResourceRecords",[])])})
                except: pass
    except: pass
    return rows

def collect_apigw(s, region):
    try:
        ag = cli(s,"apigateway",region)
        apis = ag.get_rest_apis()["items"]
        rows = []
        for api in apis:
            try: stage_count = len(ag.get_stages(restApiId=api["id"])["item"])
            except: stage_count = 0
            rows.append({"API Name":safe(api.get("name")),"API ID":safe(api.get("id")),
                         "Endpoint Type":safe(api.get("endpointConfiguration",{}).get("types",[])),
                         "Stage Count":stage_count,"Created":safe(str(api.get("createdDate",""))[:19])})
        return rows
    except: return []

def collect_apigwv2(s, region):
    try:
        ag = cli(s,"apigatewayv2",region)
        apis = ag.get_apis()["Items"]
        return [{"API Name":safe(a.get("Name")),"API ID":safe(a.get("ApiId")),
                 "Protocol":safe(a.get("ProtocolType")),"Endpoint":safe(a.get("ApiEndpoint","N/A")),
                 "CORS Config":"Yes" if a.get("CorsConfiguration") else "No",
                 "Created":safe(str(a.get("CreatedDate",""))[:19])} for a in apis]
    except: return []

def collect_vpc_peering(s, region):
    ec2 = cli(s,"ec2",region)
    peerings = ec2.describe_vpc_peering_connections()["VpcPeeringConnections"]
    return [{"Peering ID":safe(p.get("VpcPeeringConnectionId")),"Name":tag_name(p.get("Tags")),
             "Status":safe(p.get("Status",{}).get("Message")),
             "Requester VPC":safe(p.get("RequesterVpcInfo",{}).get("VpcId")),
             "Accepter VPC":safe(p.get("AccepterVpcInfo",{}).get("VpcId"))} for p in peerings]

def collect_nacl(s, region):
    ec2 = cli(s,"ec2",region)
    nacls = ec2.describe_network_acls()["NetworkAcls"]
    return [{"NACL ID":safe(n.get("NetworkAclId")),"Name":tag_name(n.get("Tags")),
             "VPC ID":safe(n.get("VpcId")),"Is Default":safe(n.get("IsDefault")),
             "Inbound Rules":len([e for e in n.get("Entries",[]) if not e.get("Egress")]),
             "Outbound Rules":len([e for e in n.get("Entries",[]) if e.get("Egress")]),
             "Associated Subnets":len(n.get("Associations",[]))} for n in nacls]

def collect_transit_gw(s, region):
    try:
        ec2 = cli(s,"ec2",region)
        tgws = ec2.describe_transit_gateways()["TransitGateways"]
        return [{"TGW ID":safe(t.get("TransitGatewayId")),"Name":tag_name(t.get("Tags")),
                 "State":safe(t.get("State")),"Owner":safe(t.get("OwnerId")),
                 "ARN":safe(t.get("TransitGatewayArn")),"Created":safe(str(t.get("CreationTime",""))[:19])} for t in tgws]
    except: return []

def collect_globalaccelerator(s, region):
    try:
        creds = s.get_credentials()
        ga = boto3.client("globalaccelerator",aws_access_key_id=creds.access_key,
                          aws_secret_access_key=creds.secret_key,region_name="us-west-2")
        accs = ga.list_accelerators()["Accelerators"]
        return [{"Name":safe(a.get("Name")),"ARN":safe(a.get("AcceleratorArn")),
                 "Status":safe(a.get("Status")),"Enabled":safe(a.get("Enabled")),
                 "DNS Name":safe(a.get("DnsName","N/A")),"Created":safe(str(a.get("CreatedTime",""))[:19])} for a in accs]
    except: return []

def collect_iam(s, region):
    iam = cli(s,"iam"); rows = []
    for page in iam.get_paginator("list_users").paginate():
        for u in page["Users"]:
            uname = u["UserName"]
            try: mfa = iam.list_mfa_devices(UserName=uname)["MFADevices"]; mfa_en = "Yes" if mfa else "No"
            except: mfa_en = "N/A"
            try: keys = iam.list_access_keys(UserName=uname)["AccessKeyMetadata"]
            except: keys = []
            rows.append({"User Name":uname,"User ID":safe(u.get("UserId")),"ARN":safe(u.get("Arn")),
                         "Path":safe(u.get("Path")),"Created":safe(str(u.get("CreateDate",""))[:19]),
                         "Password Last Used":safe(str(u.get("PasswordLastUsed","Never"))[:19]),
                         "MFA Enabled":mfa_en,"Access Key 1":safe(keys[0]["Status"]) if len(keys)>0 else "N/A",
                         "Access Key 2":safe(keys[1]["Status"]) if len(keys)>1 else "N/A"})
    return rows

def collect_iam_role(s, region):
    iam = cli(s,"iam"); rows = []
    for page in iam.get_paginator("list_roles").paginate():
        for r in page["Roles"]:
            try: attached = iam.list_attached_role_policies(RoleName=r["RoleName"])["AttachedPolicies"]; policies = ", ".join(p["PolicyName"] for p in attached)
            except: policies = "N/A"
            rows.append({"Role Name":safe(r.get("RoleName")),"Role ID":safe(r.get("RoleId")),"ARN":safe(r.get("Arn")),
                         "Description":safe(r.get("Description","N/A")),"Attached Policies":policies,
                         "Created":safe(str(r.get("CreateDate",""))[:19])})
    return rows

def collect_iam_policy(s, region):
    iam = cli(s,"iam"); rows = []
    for page in iam.get_paginator("list_policies").paginate(Scope="Local"):
        for p in page["Policies"]:
            rows.append({"Policy Name":safe(p.get("PolicyName")),"Policy ID":safe(p.get("PolicyId")),
                         "ARN":safe(p.get("Arn")),"Description":safe(p.get("Description","N/A")),
                         "Attachment Count":safe(p.get("AttachmentCount")),"Created":safe(str(p.get("CreateDate",""))[:19])})
    return rows

def collect_iam_group(s, region):
    iam = cli(s,"iam"); rows = []
    for page in iam.get_paginator("list_groups").paginate():
        for g in page["Groups"]:
            try: user_count = len(iam.get_group(GroupName=g["GroupName"])["Users"])
            except: user_count = "N/A"
            rows.append({"Group Name":safe(g.get("GroupName")),"Group ID":safe(g.get("GroupId")),
                         "ARN":safe(g.get("Arn")),"User Count":user_count,"Created":safe(str(g.get("CreateDate",""))[:19])})
    return rows

def collect_kms(s, region):
    kms = cli(s,"kms",region); rows = []
    for page in kms.get_paginator("list_keys").paginate():
        for key in page["Keys"]:
            kid = key["KeyId"]
            try: meta = kms.describe_key(KeyId=kid)["KeyMetadata"]
            except: continue
            if meta.get("KeyManager") == "AWS": continue
            try: aliases = ", ".join(a["AliasName"] for a in kms.list_aliases(KeyId=kid)["Aliases"])
            except: aliases = "N/A"
            rows.append({"Key ID":kid,"Aliases":aliases,"ARN":safe(meta.get("Arn")),
                         "Key State":safe(meta.get("KeyState")),"Key Usage":safe(meta.get("KeyUsage")),
                         "Created":safe(str(meta.get("CreationDate",""))[:19])})
    return rows

def collect_secrets(s, region):
    sm = cli(s,"secretsmanager",region); rows = []
    for page in sm.get_paginator("list_secrets").paginate():
        for sec in page["SecretList"]:
            rows.append({"Secret Name":safe(sec.get("Name")),"ARN":safe(sec.get("ARN")),
                         "Description":safe(sec.get("Description","N/A")),
                         "Rotation Enabled":safe(sec.get("RotationEnabled",False)),
                         "Last Changed":safe(str(sec.get("LastChangedDate","N/A"))[:19]),
                         "Created":safe(str(sec.get("CreatedDate",""))[:19])})
    return rows

def collect_acm(s, region):
    try:
        acm = cli(s,"acm",region); rows = []
        for page in acm.get_paginator("list_certificates").paginate():
            for cert in page["CertificateSummaryList"]:
                try: detail = acm.describe_certificate(CertificateArn=cert["CertificateArn"])["Certificate"]
                except: detail = cert
                rows.append({"Domain Name":safe(cert.get("DomainName")),"ARN":safe(cert.get("CertificateArn")),
                             "Status":safe(cert.get("Status")),"Type":safe(detail.get("Type","N/A")),
                             "Issuer":safe(detail.get("Issuer","N/A")),"In Use By":len(detail.get("InUseBy",[])),
                             "Not After":safe(str(detail.get("NotAfter","N/A"))[:19])})
        return rows
    except: return []

def collect_waf(s, region):
    rows = []
    try:
        waf = cli(s,"wafv2",region)
        for acl in waf.list_web_acls(Scope="REGIONAL").get("WebACLs",[]):
            rows.append({"Web ACL Name":safe(acl.get("Name")),"ID":safe(acl.get("Id")),
                         "ARN":safe(acl.get("ARN")),"Scope":"REGIONAL"})
    except: pass
    return rows

def collect_shield(s, region):
    try:
        creds = s.get_credentials()
        sh = boto3.client("shield",aws_access_key_id=creds.access_key,aws_secret_access_key=creds.secret_key,region_name="us-east-1")
        protections = sh.list_protections()["Protections"]
        return [{"Protection ID":safe(p.get("Id")),"Name":safe(p.get("Name")),
                 "Resource ARN":safe(p.get("ResourceArn"))} for p in protections]
    except: return []

def collect_guardduty(s, region):
    try:
        gd = cli(s,"guardduty",region)
        detectors = gd.list_detectors()["DetectorIds"]
        rows = []
        for did in detectors:
            d = gd.get_detector(DetectorId=did)
            rows.append({"Detector ID":did,"Status":safe(d.get("Status")),
                         "Finding Frequency":safe(d.get("FindingPublishingFrequency")),
                         "Created":safe(str(d.get("CreatedAt",""))[:19])})
        return rows
    except: return []

def collect_securityhub(s, region):
    try:
        sh = cli(s,"securityhub",region)
        hub = sh.describe_hub()
        standards = sh.get_enabled_standards()["StandardsSubscriptions"]
        return [{"Hub ARN":safe(hub.get("HubArn")),
                 "Subscribed At":safe(str(hub.get("SubscribedAt",""))[:19]),
                 "Standards Count":len(standards)}]
    except: return []

def collect_cognito(s, region):
    try:
        cog = cli(s,"cognito-idp",region); rows = []
        for page in cog.get_paginator("list_user_pools").paginate(MaxResults=60):
            for pool in page["UserPools"]:
                try: detail = cog.describe_user_pool(UserPoolId=pool["Id"])["UserPool"]; users = detail.get("EstimatedNumberOfUsers",0)
                except: users = "N/A"; detail = pool
                rows.append({"Pool Name":safe(pool.get("Name")),"Pool ID":safe(pool.get("Id")),
                             "User Count":users,"MFA Config":safe(detail.get("MfaConfiguration","OFF")),
                             "ARN":safe(detail.get("Arn","N/A")),"Created":safe(str(pool.get("CreationDate",""))[:19])})
        return rows
    except: return []

def collect_sns(s, region):
    sns = cli(s,"sns",region); rows = []
    for page in sns.get_paginator("list_topics").paginate():
        for t in page["Topics"]:
            arn = t["TopicArn"]; name = arn.split(":")[-1]
            try: attrs = sns.get_topic_attributes(TopicArn=arn)["Attributes"]
            except: attrs = {}
            rows.append({"Topic Name":name,"ARN":arn,"Type":"FIFO" if name.endswith(".fifo") else "Standard",
                         "Confirmed Subscriptions":safe(attrs.get("SubscriptionsConfirmed","N/A")),
                         "Pending Subscriptions":safe(attrs.get("SubscriptionsPending","N/A")),
                         "Encryption":"Enabled" if attrs.get("KmsMasterKeyId") else "Disabled"})
    return rows

def collect_sqs(s, region):
    sqs = cli(s,"sqs",region); rows = []
    for page in sqs.get_paginator("list_queues").paginate(QueueNamePrefix=""):
        for url in page.get("QueueUrls",[]):
            try: attrs = sqs.get_queue_attributes(QueueUrl=url,AttributeNames=["All"])["Attributes"]
            except: attrs = {}
            name = url.split("/")[-1]
            rows.append({"Queue Name":name,"Type":"FIFO" if name.endswith(".fifo") else "Standard",
                         "Approx Messages":safe(attrs.get("ApproximateNumberOfMessages")),
                         "Visibility Timeout":safe(attrs.get("VisibilityTimeout")),
                         "Retention (sec)":safe(attrs.get("MessageRetentionPeriod")),
                         "ARN":safe(attrs.get("QueueArn"))})
    return rows

def collect_ses(s, region):
    try:
        ses = cli(s,"ses",region)
        identities = ses.list_identities()["Identities"]
        rows = []
        for ident in identities:
            try:
                verif = ses.get_identity_verification_attributes(Identities=[ident])["VerificationAttributes"].get(ident,{})
                rows.append({"Identity":ident,"Verification Status":safe(verif.get("VerificationStatus","N/A"))})
            except: rows.append({"Identity":ident})
        return rows
    except: return []

def collect_eventbridge(s, region):
    try:
        eb = cli(s,"events",region); rows = []
        for page in eb.get_paginator("list_rules").paginate():
            for rule in page["Rules"]:
                try: target_count = len(eb.list_targets_by_rule(Rule=rule["Name"])["Targets"])
                except: target_count = 0
                rows.append({"Rule Name":safe(rule.get("Name")),"ARN":safe(rule.get("Arn")),
                             "State":safe(rule.get("State")),"Schedule":safe(rule.get("ScheduleExpression","N/A")),
                             "Event Bus":safe(rule.get("EventBusName","default")),"Targets":target_count})
        return rows
    except: return []

def collect_kinesis(s, region):
    try:
        k = cli(s,"kinesis",region)
        streams = k.list_streams()["StreamNames"]
        rows = []
        for name in streams:
            try:
                detail = k.describe_stream_summary(StreamName=name)["StreamDescriptionSummary"]
                rows.append({"Stream Name":safe(detail.get("StreamName")),"ARN":safe(detail.get("StreamARN")),
                             "Status":safe(detail.get("StreamStatus")),"Shard Count":safe(detail.get("OpenShardCount")),
                             "Retention (hours)":safe(detail.get("RetentionPeriodHours")),"Encryption":safe(detail.get("EncryptionType","NONE"))})
            except: rows.append({"Stream Name":name})
        return rows
    except: return []

def collect_lambda(s, region):
    lm = cli(s,"lambda",region); rows = []
    for page in lm.get_paginator("list_functions").paginate():
        for fn in page["Functions"]:
            rows.append({"Function Name":safe(fn.get("FunctionName")),"ARN":safe(fn.get("FunctionArn")),
                         "Runtime":safe(fn.get("Runtime")),"Handler":safe(fn.get("Handler")),
                         "Memory (MB)":safe(fn.get("MemorySize")),"Timeout (sec)":safe(fn.get("Timeout")),
                         "Code Size (bytes)":safe(fn.get("CodeSize")),"State":safe(fn.get("State","Active")),
                         "Role":safe(fn.get("Role")),"Last Modified":safe(fn.get("LastModified","")[:19]),
                         "Architecture":safe(fn.get("Architectures",["x86_64"]))})
    return rows

def collect_cloudwatch(s, region):
    cw = cli(s,"cloudwatch",region); rows = []
    for page in cw.get_paginator("describe_alarms").paginate():
        for alarm in page.get("MetricAlarms",[]):
            rows.append({"Alarm Name":safe(alarm.get("AlarmName")),"State":safe(alarm.get("StateValue")),
                         "Metric":safe(alarm.get("MetricName")),"Namespace":safe(alarm.get("Namespace")),
                         "Threshold":safe(alarm.get("Threshold")),"Actions Enabled":safe(alarm.get("ActionsEnabled")),
                         "ARN":safe(alarm.get("AlarmArn"))})
    return rows

def collect_cloudwatch_log(s, region):
    try:
        logs = cli(s,"logs",region); rows = []
        for page in logs.get_paginator("describe_log_groups").paginate():
            for lg in page["logGroups"]:
                rows.append({"Log Group Name":safe(lg.get("logGroupName")),
                             "Retention (days)":safe(lg.get("retentionInDays","Never Expire")),
                             "Stored Bytes":safe(lg.get("storedBytes")),
                             "ARN":safe(lg.get("arn"))})
        return rows
    except: return []

def collect_cloudtrail(s, region):
    try:
        ct = cli(s,"cloudtrail",region)
        trails = ct.describe_trails(includeShadowTrails=False)["trailList"]
        rows = []
        for t in trails:
            try: logging = safe(ct.get_trail_status(Name=t["TrailARN"]).get("IsLogging"))
            except: logging = "N/A"
            rows.append({"Trail Name":safe(t.get("Name")),"ARN":safe(t.get("TrailARN")),
                         "S3 Bucket":safe(t.get("S3BucketName")),"Multi-Region":safe(t.get("IsMultiRegionTrail")),
                         "Home Region":safe(t.get("HomeRegion")),"Log Validation":safe(t.get("LogFileValidationEnabled")),
                         "Is Logging":logging})
        return rows
    except: return []

def collect_cloudformation(s, region):
    try:
        cf = cli(s,"cloudformation",region); rows = []
        for page in cf.get_paginator("list_stacks").paginate(StackStatusFilter=["CREATE_COMPLETE","UPDATE_COMPLETE","ROLLBACK_COMPLETE","UPDATE_ROLLBACK_COMPLETE"]):
            for stack in page["StackSummaries"]:
                rows.append({"Stack Name":safe(stack.get("StackName")),"Status":safe(stack.get("StackStatus")),
                             "Description":safe(stack.get("TemplateDescription","N/A")),
                             "Drift Status":safe(stack.get("DriftInformation",{}).get("StackDriftStatus","N/A")),
                             "Created":safe(str(stack.get("CreationTime",""))[:19]),"Updated":safe(str(stack.get("LastUpdatedTime","N/A"))[:19])})
        return rows
    except: return []

def collect_ssm(s, region):
    try:
        ssm = cli(s,"ssm",region); rows = []
        for page in ssm.get_paginator("describe_instance_information").paginate():
            for i in page["InstanceInformationList"]:
                rows.append({"Instance ID":safe(i.get("InstanceId")),"Ping Status":safe(i.get("PingStatus")),
                             "Agent Version":safe(i.get("AgentVersion")),"Platform":safe(i.get("PlatformType")),
                             "Platform Name":safe(i.get("PlatformName")),"Computer Name":safe(i.get("ComputerName","N/A"))})
        return rows
    except: return []

def collect_ssm_param(s, region):
    try:
        ssm = cli(s,"ssm",region); rows = []
        for page in ssm.get_paginator("describe_parameters").paginate():
            for p in page["Parameters"]:
                rows.append({"Parameter Name":safe(p.get("Name")),"Type":safe(p.get("Type")),
                             "Description":safe(p.get("Description","N/A")),"Version":safe(p.get("Version")),
                             "Tier":safe(p.get("Tier","Standard")),"Last Modified":safe(str(p.get("LastModifiedDate",""))[:19])})
        return rows
    except: return []

def collect_stepfunctions(s, region):
    try:
        sf = cli(s,"stepfunctions",region)
        machines = sf.list_state_machines()["stateMachines"]
        rows = []
        for m in machines:
            try: detail = sf.describe_state_machine(stateMachineArn=m["stateMachineArn"])
            except: detail = m
            rows.append({"Name":safe(m.get("name")),"ARN":safe(m.get("stateMachineArn")),
                         "Type":safe(detail.get("type","N/A")),"Status":safe(detail.get("status","ACTIVE")),
                         "Created":safe(str(m.get("creationDate",""))[:19])})
        return rows
    except: return []

def collect_eks(s, region):
    eks = cli(s,"eks",region); rows = []
    for cname in eks.list_clusters().get("clusters",[]):
        try:
            c = eks.describe_cluster(name=cname)["cluster"]
            rc = c.get("resourcesVpcConfig",{})
            rows.append({"Cluster Name":safe(c.get("name")),"ARN":safe(c.get("arn")),
                         "Status":safe(c.get("status")),"K8s Version":safe(c.get("version")),
                         "VPC ID":safe(rc.get("vpcId")),"Endpoint Public":safe(rc.get("endpointPublicAccess")),
                         "Endpoint Private":safe(rc.get("endpointPrivateAccess")),
                         "Created":safe(str(c.get("createdAt",""))[:19])})
        except Exception as e:
            rows.append({"Cluster Name":cname,"Error":str(e)})
    return rows

def collect_ecs(s, region):
    ecs = cli(s,"ecs",region)
    arns = ecs.list_clusters().get("clusterArns",[])
    if not arns: return []
    clusters = ecs.describe_clusters(clusters=arns,include=["STATISTICS","SETTINGS"])["clusters"]
    rows = []
    for c in clusters:
        stats = {st["name"]:st["value"] for st in c.get("statistics",[])}
        rows.append({"Cluster Name":safe(c.get("clusterName")),"ARN":safe(c.get("clusterArn")),
                     "Status":safe(c.get("status")),"Running Tasks":safe(stats.get("runningTasksCount","0")),
                     "Active Services":safe(stats.get("activeServicesCount","0")),
                     "Instances":safe(stats.get("registeredContainerInstancesCount","0"))})
    return rows

def collect_ecr(s, region):
    ecr = cli(s,"ecr",region); rows = []
    for page in ecr.get_paginator("describe_repositories").paginate():
        for repo in page["repositories"]:
            try: img_count = len(ecr.list_images(repositoryName=repo["repositoryName"])["imageIds"])
            except: img_count = "N/A"
            rows.append({"Repository Name":safe(repo.get("repositoryName")),"ARN":safe(repo.get("repositoryArn")),
                         "URI":safe(repo.get("repositoryUri")),"Image Count":img_count,
                         "Tag Mutability":safe(repo.get("imageTagMutability")),
                         "Scan on Push":safe(repo.get("imageScanningConfiguration",{}).get("scanOnPush",False)),
                         "Created":safe(str(repo.get("createdAt",""))[:19])})
    return rows

def collect_sagemaker(s, region):
    try:
        sm = cli(s,"sagemaker",region); rows = []
        for page in sm.get_paginator("list_notebook_instances").paginate():
            for nb in page["NotebookInstances"]:
                rows.append({"Name":safe(nb.get("NotebookInstanceName")),"Status":safe(nb.get("NotebookInstanceStatus")),
                             "Instance Type":safe(nb.get("InstanceType")),"ARN":safe(nb.get("NotebookInstanceArn")),
                             "Created":safe(str(nb.get("CreationTime",""))[:19])})
        return rows
    except: return []

def collect_glue(s, region):
    try:
        glue = cli(s,"glue",region); rows = []
        try:
            for page in glue.get_paginator("get_databases").paginate():
                for db in page["DatabaseList"]:
                    rows.append({"Resource":"Database","Name":safe(db.get("Name")),
                                 "Description":safe(db.get("Description","N/A")),"Created":safe(str(db.get("CreateTime",""))[:19])})
        except: pass
        try:
            for page in glue.get_paginator("get_crawlers").paginate():
                for c in page["Crawlers"]:
                    rows.append({"Resource":"Crawler","Name":safe(c.get("Name")),"Status":safe(c.get("State")),
                                 "Role":safe(c.get("Role","N/A"))})
        except: pass
        return rows
    except: return []

def collect_emr(s, region):
    try:
        emr = cli(s,"emr",region)
        clusters = emr.list_clusters()["Clusters"]
        return [{"Cluster ID":safe(c.get("Id")),"Name":safe(c.get("Name")),
                 "Status":safe(c.get("Status",{}).get("State")),
                 "Created":safe(str(c.get("Status",{}).get("Timeline",{}).get("CreationDateTime",""))[:19])} for c in clusters]
    except: return []

def collect_opensearch(s, region):
    try:
        os = cli(s,"opensearch",region)
        domains = os.list_domain_names()["DomainNames"]
        rows = []
        for d in domains:
            try:
                detail = os.describe_domain(DomainName=d["DomainName"])["DomainStatus"]
                rows.append({"Domain Name":safe(detail.get("DomainName")),"ARN":safe(detail.get("ARN")),
                             "Engine Version":safe(detail.get("EngineVersion","N/A")),
                             "Endpoint":safe(detail.get("Endpoint","N/A")),
                             "Instance Type":safe(detail.get("ClusterConfig",{}).get("InstanceType","N/A")),
                             "Encrypted":safe(detail.get("EncryptionAtRestOptions",{}).get("Enabled","N/A"))})
            except: rows.append({"Domain Name":d["DomainName"],"Error":"Failed"})
        return rows
    except: return []

def collect_elasticbeanstalk(s, region):
    try:
        eb = cli(s,"elasticbeanstalk",region)
        envs = eb.describe_environments()["Environments"]
        return [{"Env Name":safe(e.get("EnvironmentName")),"App Name":safe(e.get("ApplicationName")),
                 "Status":safe(e.get("Status")),"Health":safe(e.get("Health")),
                 "CNAME":safe(e.get("CNAME","N/A")),"Solution Stack":safe(e.get("SolutionStackName","N/A")),
                 "Created":safe(str(e.get("DateCreated",""))[:19])} for e in envs]
    except: return []

def collect_codecommit(s, region):
    try:
        cc = cli(s,"codecommit",region)
        repos = cc.list_repositories()["repositories"]
        rows = []
        for repo in repos:
            try:
                detail = cc.get_repository(repositoryName=repo["repositoryName"])["repositoryMetadata"]
                rows.append({"Repository Name":safe(detail.get("repositoryName")),"Default Branch":safe(detail.get("defaultBranch","N/A")),
                             "ARN":safe(detail.get("Arn")),"Created":safe(str(detail.get("creationDate",""))[:19])})
            except: rows.append({"Repository Name":repo["repositoryName"]})
        return rows
    except: return []

def collect_codebuild(s, region):
    try:
        cb = cli(s,"codebuild",region)
        projects = cb.list_projects()["projects"]
        if not projects: return []
        details = cb.batch_get_projects(names=projects[:20])["projects"]
        return [{"Project Name":safe(p.get("name")),"ARN":safe(p.get("arn")),
                 "Source Type":safe(p.get("source",{}).get("type","N/A")),
                 "Build Image":safe(p.get("environment",{}).get("image","N/A")),
                 "Service Role":safe(p.get("serviceRole","N/A")),"Created":safe(str(p.get("created",""))[:19])} for p in details]
    except: return []

def collect_codepipeline(s, region):
    try:
        cp = cli(s,"codepipeline",region)
        pipelines = cp.list_pipelines()["pipelines"]
        return [{"Pipeline Name":safe(p.get("name")),"Version":safe(p.get("version")),
                 "Created":safe(str(p.get("created",""))[:19]),"Updated":safe(str(p.get("updated",""))[:19])} for p in pipelines]
    except: return []

def collect_xray(s, region):
    try:
        xr = cli(s,"xray",region)
        groups = xr.get_groups()["Groups"]
        return [{"Group Name":safe(g.get("GroupName")),"Group ARN":safe(g.get("GroupARN")),
                 "Filter":safe(g.get("FilterExpression","N/A"))} for g in groups]
    except: return []

def collect_config(s, region):
    try:
        cfg = cli(s,"config",region)
        recorders = cfg.describe_configuration_recorders()["ConfigurationRecorders"]
        statuses = {st["name"]:st for st in cfg.describe_configuration_recorder_status()["ConfigurationRecordersStatus"]}
        return [{"Recorder Name":safe(r.get("name")),"Role ARN":safe(r.get("roleARN","N/A")),
                 "All Supported":safe(r.get("recordingGroup",{}).get("allSupported")),
                 "Recording":safe(statuses.get(r["name"],{}).get("recording","N/A"))} for r in recorders]
    except: return []

def collect_organizations(s, region):
    try:
        creds = s.get_credentials()
        org = boto3.client("organizations",aws_access_key_id=creds.access_key,aws_secret_access_key=creds.secret_key,region_name="us-east-1")
        accounts = []
        for page in org.get_paginator("list_accounts").paginate():
            accounts.extend(page["Accounts"])
        return [{"Account ID":safe(a.get("Id")),"Name":safe(a.get("Name")),
                 "Email":safe(a.get("Email")),"Status":safe(a.get("Status")),
                 "Joined Method":safe(a.get("JoinedMethod")),"Joined":safe(str(a.get("JoinedTimestamp",""))[:19])} for a in accounts]
    except: return []

def collect_mq(s, region):
    try:
        mq = cli(s,"mq",region)
        brokers = mq.list_brokers()["BrokerSummaries"]
        return [{"Broker Name":safe(b.get("BrokerName")),"Broker ID":safe(b.get("BrokerId")),
                 "State":safe(b.get("BrokerState")),"Engine Type":safe(b.get("EngineType")),
                 "Host Type":safe(b.get("HostInstanceType")),"ARN":safe(b.get("BrokerArn")),
                 "Created":safe(str(b.get("Created",""))[:19])} for b in brokers]
    except: return []

def collect_lambda_layer(s, region):
    try:
        lm = cli(s,"lambda",region); rows = []
        for page in lm.get_paginator("list_layers").paginate():
            for layer in page["Layers"]:
                lv = layer.get("LatestMatchingVersion",{})
                rows.append({"Layer Name":safe(layer.get("LayerName")),"Layer ARN":safe(layer.get("LayerArn")),
                             "Latest Version":safe(lv.get("Version")),"Runtimes":safe(lv.get("CompatibleRuntimes",[])),
                             "Created":safe(str(lv.get("CreatedDate",""))[:19])})
        return rows
    except: return []

def collect_dax(s, region):
    try:
        dax = cli(s,"dax",region)
        clusters = dax.describe_clusters()["Clusters"]
        return [{"Cluster Name":safe(c.get("ClusterName")),"Status":safe(c.get("Status")),
                 "Node Type":safe(c.get("NodeType")),"Total Nodes":safe(c.get("TotalNodes")),
                 "ARN":safe(c.get("ClusterArn"))} for c in clusters]
    except: return []

def collect_pinpoint(s, region):
    try:
        pp = cli(s,"pinpoint",region)
        apps = pp.get_apps()["ApplicationsResponse"]["Item"]
        return [{"App Name":safe(a.get("Name")),"App ID":safe(a.get("Id")),
                 "ARN":safe(a.get("Arn")),"Created":safe(str(a.get("CreationDate",""))[:19])} for a in apps]
    except: return []

def collect_appsync(s, region):
    try:
        appsync = cli(s,"appsync",region)
        apis = appsync.list_graphql_apis()["graphqlApis"]
        return [{"API Name":safe(a.get("name")),"API ID":safe(a.get("apiId")),
                 "ARN":safe(a.get("arn")),"Auth Type":safe(a.get("authenticationType"))} for a in apis]
    except: return []

def collect_amplify(s, region):
    try:
        amp = cli(s,"amplify",region)
        apps = amp.list_apps()["apps"]
        return [{"App Name":safe(a.get("name")),"App ID":safe(a.get("appId")),
                 "ARN":safe(a.get("appArn")),"Platform":safe(a.get("platform","N/A")),
                 "Created":safe(str(a.get("createTime",""))[:19])} for a in apps]
    except: return []

def collect_transfer(s, region):
    try:
        tf = cli(s,"transfer",region)
        servers = tf.list_servers()["Servers"]
        return [{"Server ID":safe(sv.get("ServerId")),"ARN":safe(sv.get("Arn")),
                 "Domain":safe(sv.get("Domain","N/A")),"State":safe(sv.get("State"))} for sv in servers]
    except: return []

def collect_workspaces(s, region):
    try:
        ws = cli(s,"workspaces",region); rows = []
        for page in ws.get_paginator("describe_workspaces").paginate():
            for w in page["Workspaces"]:
                rows.append({"Workspace ID":safe(w.get("WorkspaceId")),"Username":safe(w.get("UserName")),
                             "State":safe(w.get("State")),"Compute":safe(w.get("WorkspaceProperties",{}).get("ComputeTypeName","N/A")),
                             "IP Address":safe(w.get("IpAddress","N/A"))})
        return rows
    except: return []

def collect_storagegateway(s, region):
    try:
        sg = cli(s,"storagegateway",region)
        gateways = sg.list_gateways()["Gateways"]
        return [{"Gateway Name":safe(g.get("GatewayName")),"ID":safe(g.get("GatewayId")),
                 "Type":safe(g.get("GatewayType")),"State":safe(g.get("GatewayOperationalState"))} for g in gateways]
    except: return []

def collect_kinesis_firehose(s, region):
    try:
        kf = cli(s,"firehose",region)
        streams = kf.list_delivery_streams()["DeliveryStreamNames"]
        rows = []
        for name in streams:
            try:
                detail = kf.describe_delivery_stream(DeliveryStreamName=name)["DeliveryStreamDescription"]
                rows.append({"Stream Name":safe(detail.get("DeliveryStreamName")),"ARN":safe(detail.get("DeliveryStreamARN")),
                             "Status":safe(detail.get("DeliveryStreamStatus")),"Type":safe(detail.get("DeliveryStreamType")),
                             "Created":safe(str(detail.get("CreateTimestamp",""))[:19])})
            except: rows.append({"Stream Name":name})
        return rows
    except: return []

def collect_msk(s, region):
    try:
        msk = cli(s,"kafka",region)
        clusters = msk.list_clusters()["ClusterInfoList"]
        return [{"Cluster Name":safe(c.get("ClusterName")),"ARN":safe(c.get("ClusterArn")),
                 "State":safe(c.get("State")),"Kafka Version":safe(c.get("CurrentBrokerSoftwareInfo",{}).get("KafkaVersion","N/A")),
                 "Brokers":safe(c.get("NumberOfBrokerNodes")),"Created":safe(str(c.get("CreationTime",""))[:19])} for c in clusters]
    except: return []

COLLECTORS = {
    "ec2": collect_ec2, "ec2_ami": collect_ec2_ami, "ec2_sg": collect_ec2_sg,
    "ec2_keypair": collect_ec2_keypair, "ec2_eip": collect_ec2_eip,
    "ec2_snapshot": collect_ec2_snapshot, "ec2_volume": collect_ec2_volume,
    "autoscaling": collect_autoscaling, "lightsail": collect_lightsail, "batch": collect_batch,
    "s3": collect_s3, "efs": collect_efs, "fsx": collect_fsx,
    "backup": collect_backup, "storagegateway": collect_storagegateway,
    "rds": collect_rds, "rds_snapshot": collect_rds_snapshot,
    "dynamodb": collect_dynamodb, "elasticache": collect_elasticache,
    "redshift": collect_redshift, "docdb": collect_docdb, "neptune": collect_neptune,
    "vpc": collect_vpc, "subnet": collect_subnet, "routetable": collect_routetable,
    "igw": collect_igw, "nat": collect_nat, "vpn": collect_vpn,
    "directconnect": collect_directconnect, "elb": collect_elb,
    "cloudfront": collect_cloudfront, "route53": collect_route53,
    "route53_record": collect_route53_record, "apigw": collect_apigw,
    "apigwv2": collect_apigwv2, "vpc_peering": collect_vpc_peering,
    "nacl": collect_nacl, "transit_gw": collect_transit_gw,
    "globalaccelerator": collect_globalaccelerator,
    "iam": collect_iam, "iam_role": collect_iam_role, "iam_policy": collect_iam_policy,
    "iam_group": collect_iam_group, "kms": collect_kms, "secrets": collect_secrets,
    "acm": collect_acm, "waf": collect_waf, "shield": collect_shield,
    "guardduty": collect_guardduty, "securityhub": collect_securityhub, "cognito": collect_cognito,
    "sns": collect_sns, "sqs": collect_sqs, "ses": collect_ses,
    "eventbridge": collect_eventbridge, "mq": collect_mq,
    "kinesis": collect_kinesis, "kinesis_firehose": collect_kinesis_firehose,
    "glue": collect_glue, "emr": collect_emr, "opensearch": collect_opensearch, "msk": collect_msk,
    "lambda": collect_lambda, "lambda_layer": collect_lambda_layer,
    "eks": collect_eks, "ecs": collect_ecs, "ecr": collect_ecr,
    "cloudwatch": collect_cloudwatch, "cloudwatch_log": collect_cloudwatch_log,
    "cloudtrail": collect_cloudtrail, "cloudformation": collect_cloudformation,
    "config": collect_config, "ssm": collect_ssm, "ssm_param": collect_ssm_param,
    "stepfunctions": collect_stepfunctions, "xray": collect_xray,
    "sagemaker": collect_sagemaker, "codecommit": collect_codecommit,
    "codebuild": collect_codebuild, "codepipeline": collect_codepipeline,
    "elasticbeanstalk": collect_elasticbeanstalk, "organizations": collect_organizations,
    "dax": collect_dax, "pinpoint": collect_pinpoint, "appsync": collect_appsync,
    "amplify": collect_amplify, "transfer": collect_transfer, "workspaces": collect_workspaces,
}

if __name__ == "__main__":
    import webbrowser, threading, time
    print("=" * 60)
    print("  AWS Cloud Management Suite — Shellkode")
    print("  Modules: Inventory | Cost | Security | Monthly Report")
    print("  Open: http://localhost:8080")
    print("  Press Ctrl+C to stop")
    print("=" * 60)
    def open_browser():
        time.sleep(1.2)
        webbrowser.open("http://localhost:8080")
    threading.Thread(target=open_browser, daemon=True).start()
    app.run(host="0.0.0.0", port=8080, debug=False)
